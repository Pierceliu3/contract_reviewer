import re
import streamlit as st
from docx import Document
import pypdf
from datetime import datetime
from rapidfuzz import fuzz
from langchain_ollama import ChatOllama
import json
from openai import OpenAI
import os
import hmac
from report_generator import get_download_button_data 




ZHIPU_API_KEY_1 = st.secrets.get("ZHIPU_API_KEY_1")
ZHIPU_API_KEY_2 = st.secrets.get("ZHIPU_API_KEY_2") 
ZHIPU_API_KEY_3 = st.secrets.get("ZHIPU_API_KEY_3")
ZHIPU_API_KEY_4 = st.secrets.get("ZHIPU_API_KEY_4")


# ============================================
# 合规检查项列表（用于智能分析）
# ============================================
QUALIFICATION_STANDARDS = {
    "地基基础工程专业承包": {
        "一级": "可承担各类地基基础工程的施工",
        "二级": "高度 100 米以下工业、民用建筑工程和高度 120 米以下构筑物的地基基础工程；深度不超过 24 米的刚性桩复合地基处理和深度不超过 10 米的其它地基处理工程；单桩承受设计荷载 5000 千牛以下的桩基础工程；开挖深度不超过 15 米的基坑围护工程。",
        "三级": "高度 50 米以下工业、民用建筑工程和高度 70 米以下构筑物的地基基础工程；深度不超过 18 米的刚性桩复合地基处理或深度不超过 8 米的其它地基处理工程；单桩承受设计荷载 3000 千牛以下的桩基础工程；开挖深度不超过 12 米的基坑围护工程"
    },
    "起重设备安装工程专业承包": {
        "一级": "可承担塔式起重机、各类施工升降机和门式起重机的安装与拆卸。",
        "二级": "可承担 3150 千牛·米以下塔式起重机、各类施工升降机和门式起重机的安装与拆卸。",
        "三级": "可承担 800 千牛·米以下塔式起重机、各类施工升降机和门式起重机的安装与拆卸。"
    },
    "预拌混凝土专业承包": {
        "一级": "",
        "二级": "",
        "三级": ""
    },
    "电子与智能化工程专业承包": {
        "一级": "可承担各类型电子工程、建筑智能化工程施工。",
        "二级": "可承担单项合同额 2500 万元以下的电子工业制造设备安装工程和电子工业环境工程、单项合同额 1500 万元以下的电子系统工程和建筑智能化工程施工。"
    },
    "消防设施工程专业承包": {
        "一级": "可承担各类型消防设施工程的施工。",
        "二级": "可承担单体建筑面积 5 万平方米以下的下列消防设施工程的施工中的一类高层民用建筑以外的民用建筑和火灾危险性丙类以下的厂房、仓库、储罐、堆场。"
    },
    "防水防腐保温工程专业承包": {
        "一级": "可承担各类建筑防水、防腐保温工程的施工。",
        "二级": "可承担单项合同额 300 万元以下建筑防水工程的施工，单项合同额 600 万元以下的各类防腐保温工程的施工。"
    },
    "桥梁工程专业承包": {
        "一级": "可承担各类桥梁工程的施工。",
        "二级": "可承担单跨 150 米以下、单座桥梁总长 1000 米以下桥梁工程的施工。",
        "三级": "可承担单跨 50 米以下、单座桥梁总长 120 米以下桥梁工程的施工。"
    },
    "隧道工程专业承包": {
        "一级": "可承担各类隧道工程的施工。",
        "二级": "可承担断面 60 平方米以下且单洞长度 1000 米以下的隧道工程施工。",
        "三级": "可承担断面 40 平方米以下且单洞长度 500 米以下的隧道工程施工。"
    },
    "钢结构工程专业承包": {
        "一级": "（1）钢结构高度 60 米以上； （2）钢结构单跨跨度 30 米以上； （3）网壳、网架结构短边边跨跨度 50 米以上； （4）单体钢结构工程钢结构总重量 4000 吨以上； （5）单体建筑面积 30000 平方米以上。",
        "二级": "（1）钢结构高度 100 米以下； （2）钢结构单跨跨度 36 米以下； （3）网壳、网架结构短边边跨跨度 75 米以下； （4）单体钢结构工程钢结构总重量 6000 吨以下； （5）单体建筑面积 35000 平方米以下。",
        "三级": "（1）钢结构高度 60 米以下； （2）钢结构单跨跨度 30 米以下； （3）网壳、网架结构短边边跨跨度 33 米以下； （4）单体钢结构工程钢结构总重量 3000 吨以下； （5）单体建筑面积 15000 平方米以下。"
    },
    "模板脚手架专业承包": {
        "一级": "",
        "二级": "",
        "三级": ""
    },
    "建筑装修装饰工程专业承包": {
        "一级": "可承担各类建筑装修装饰工程，以及与装修工程直接配套的其他工程的施工。",
        "二级": "可承担单项合同额 2000 万元以下的建筑装修装饰工程，以及与装修工程直接配套的其他工程的施工。"
    },
    "建筑机电安装工程专业承包": {
        "一级": "可承担各类建筑工程项目的设备、线路、管道的安装，35 千伏以下变配电站工程，非标准钢结构件的制作、安装。",
        "二级": "可承担单项合同额 2000 万元以下的各类建筑工程项目的设备、线路、管道的安装，10 千伏以下变配电站工程，非标准钢结构件的制作、安装。",
        "三级": "可承担单项合同额 1000 万元以下的各类建筑工程项目的设备、线路、管道的安装，非标准钢结构件的制作、安装。"
    },
    "建筑幕墙工程专业承包": {
        "一级": "可承担各类型的建筑幕墙工程的施工。",
        "二级": "可承担单体建筑工程幕墙面积 8000 平方米以下建筑幕墙工程的施工。"
    },
    "古建筑工程专业承包": {
        "一级": "可承担各类仿古建筑、古建筑修缮工程的施工。",
        "二级": "可承担建筑面积 800 平方米以下的单体仿古建筑工程，国家级 200 平方米以下重点文物保护单位的古建筑修缮工程的施工。",
        "三级": "可承担建筑面积 400 平方米以下的单体仿古建筑工程，省级 100 平方米以下重点文物保护单位的古建筑修缮工程的施工。"
    },
    "城市及道路照明工程专业承包": {
        "一级": "可承担各类城市与道路照明工程的施工。",
        "二级": "可承担单项合同额不超过 1200 万元的城市与道路照明工程的施工。",
        "三级": "可承担单项合同额不超过 600 万元的城市与道路照明工程的施工。"
    },
    "公路路面工程专业承包": {
        "一级": "可承担各级公路路面工程的施工。",
        "二级": "可承担一级以下公路路面工程的施工。",
        "三级": "可承担二级以下公路路面工程的施工。"
    },
    "公路路基工程专业承包": {
        "一级": "可承担各级公路的路基、中小桥涵、防护及排水、软基处理工程的施工。",
        "二级": "可承担一级标准以下公路的路基、中小桥涵、防护及排水、软基处理工程的施工。",
        "三级": "可承担二级标准以下公路的路基、中小桥涵、防护及排水、软基处理工程的施工。"
    },
    "公路交通工程专业承包": {
        "一级": "可承担各级公路标志、标线、护栏、隔离栅、防眩板等公路安全设施工程的施工及安装。",
        "二级": "可承担一级以下公路标志、标线、护栏、隔离栅、防眩板等公路安全设施工程的施工及安装。",
        "三级": ""
    },
    "铁路电务工程专业承包": {
        "一级": "可承担各类铁路通信、信号及电力工程施工。",
        "二级": "可承担 100 公里以下Ⅰ、Ⅱ、Ⅲ、Ⅳ级铁路通信、信号及电力工程施工。",
        "三级": "可承担 50 公里以下Ⅱ、Ⅲ、Ⅳ级铁路通信、信号及电力工程施工。"
    },
    "铁路铺轨架梁工程专业承包": {
        "一级": "可承担各类大中型铁路铺轨架梁工程施工。",
        "二级": "可承担 50 公里以下Ⅰ级铁路、100 公里以下既有线改造以及Ⅱ、Ⅲ、Ⅳ级铁路铺轨架梁工程施工。"
    },
    "铁路电气化工程专业承包": {
        "一级": "可承担各类铁路电气化工程的施工。",
        "二级": "可承担 100 公里以下Ⅰ级铁路和Ⅱ、Ⅲ、Ⅳ级铁路电气化工程施工。",
        "三级": "可承担铁路站线改造和 50 公里以下Ⅱ、Ⅲ、Ⅳ级铁路电气化工程施工。"
    },
    "机场场道工程专业承包": {
        "一级": "可承担各类机场场道工程的施工。",
        "二级": "可承担飞行区指标为 4E 以上，单项合同额在 2000 万以下技术不复杂的飞行区场道工程的施工；或飞行区指标为 4D，单项合同额在 4000 万以下的飞行区场道工程的施工；或飞行区指标为 4C 以下，单项合同额在 6000 万以下的飞行区场道工程的施工；各类场道维修工程。"
    },
    "民航空管工程及机场弱电系统工程专业承包": {
        "一级": "可承担各类民航空管工程和机场弱电系统工程的施工。",
        "二级": "可承担单项合同额 2000 万元以下的民航空管工程和单项合同额 2500 万元以下的机场弱电系统工程的施工。"
    },
    "机场目视助航工程专业承包": {
        "一级": "可承担各类机场目视助航工程的施工。",
        "二级": "可承担飞行区指标为 4E 以上，单项合同额 500 万元以下的目视助航工程；或飞行区指标为 4D 以下的目视助航工程的施工。"
    },
    "港口与海岸工程专业承包": {
        "一级": "可承担各类港口与海岸工程的施工，包括码头、防波堤、护岸、围堰、堆场道路及陆域构筑物、筒仓、船坞、船台、滑道、水下地基及基础、土石方、海上灯塔、航标与警戒标志、栈桥、人工岛及平台、海上风电、海岸与近海等工程。",
        "二级": "可承担下列港口与海岸工程的施工，包括沿海 5 万吨级及内河 5000 吨级以下码头、水深小于 7 米的防波堤、5 万吨级以下船坞船台及滑道工程、1200 米以下围堤护岸工程，以及相应的堆场道路及陆域构筑物、筒仓、水下地基及基础、土石方、海上灯塔、航标与警戒标志、栈桥、人工岛及平台、海岸与近海等工程。",
        "三级": "可承担下列港口与海岸工程的施工，包括沿海 1 万吨级及内河 3000 吨级以下码头、水深小于 4 米的防波堤、1 万吨级以下船坞船台及滑道工程、800 米以下围堤护岸工程，以及相应的堆场道路及陆域构筑物、水下地基及基础、土石方、航标与警戒标志、栈桥、海岸与近海等工程。"
    },
    "航道工程专业承包": {
        "一级": "可承担各类航道工程的施工，包括河海湖航道整治（含堤、坝、护岸）、测量、航标与渠化工程，疏浚与吹填造地（含围堰），水下清障、开挖、清淤、炸礁清礁等工程。",
        "二级": "可承担沿海 5 万吨级和内河 1000 吨级以下航道工程、600 万立方米以下疏浚工程或陆域吹填工程、6 万立方米以下水下炸礁清礁工程，以及相应的测量、航标与渠化工程、水下清障、开挖、清淤等工程的施工。",
        "三级": "可承担沿海 2 万吨级和内河 500 吨级以下航道工程、300 万立方米以下疏浚工程或陆域吹填工程、4 万立方米以下水下炸礁清礁工程，以及相应的测量、航标与渠化工程、水下清障、开挖、清淤等工程的施工。"
    },
    "通航建筑物工程专业承包": {
        "一级": "可承担各类船闸、升船机等通航建筑物工程的施工。",
        "二级": "可承担 1000 吨级以下船闸或 300 吨级以下升船机等通航建筑物工程的施工。",
        "三级": "可承担 300 吨级以下船闸或 50 吨级以下升船机等通航建筑物工程的施工。"
    },
    "港航设备安装及水上交管工程专业承包": {
        "一级": "可承担各类港口装卸设备安装及配套工程的施工，各类船闸、升船机、航电枢纽设备安装工程的施工，各类水上交通管制工程的施工。",
        "二级": "可承担沿海 5 万吨级和内河 5000 吨级以下散货（含油、气）、杂货和集装箱码头成套装卸设备安装工程，1000 吨级以下船闸或 300 吨级以下升船机设备安装工程施工，单项合同额 1000 万元以下的各类水上交通管制工程的施工。"
    },
    "水工金属结构制作与安装工程专业承包": {
        "一级": "可承担各类压力钢管、闸门、拦污栅等水工金属结构工程的制作、安装及启闭机的安装。",
        "二级": "可承担大型以下压力钢管、闸门、拦污栅等水工金属结构工程的制作、安装及启闭机的安装。",
        "三级": "可承担中型以下压力钢管、闸门、拦污栅等水工金属结构工程的制作、安装及启闭机的安装。"
    },
    "水利水电机电安装工程专业承包": {
        "一级": "可承担各类水电站、泵站主机（各类水轮发电机组、水泵机组）及其附属设备和水电（泵）站电气设备的安装工程。",
        "二级": "可承担单机容量 100MW 以下的水电站、单机容量 1000KW 以下的泵站主机及其附属设备和水电（泵）站电气设备的安装工程。",
        "三级": "可承担单机容量 25MW 以下的水电站、单机容量 500KW 以下的泵站主机及其附属设备和水电(泵)站电气设备的安装工程。"
    },
    "河湖整治工程专业承包": {
        "一级": "可承担各类河道、水库、湖泊以及沿海相应工程的河势控导、险工处理、疏浚与吹填、清淤、填塘固基工程的施工。",
        "二级": "可承担堤防工程级别 2 级以下堤防相应的河道、湖泊的河势控导、险工处理、疏浚与吹填、填塘固基工程的施工。",
        "三级": "可承担堤防工程级别 3 级以下堤防相应的河湖疏浚整治工程及吹填工程的施工。"
    },
    "输变电工程专业承包": {
        "一级": "可承担各种电压等级的送电线路和变电站工程的施工。",
        "二级": "可承担 220 千伏以下电压等级的送电线路和变电站工程的施工。",
        "三级": "可承担 110 千伏以下电压等级的送电线路和变电站工程的施工。"
    },
    "核工程专业承包": {
        "一级": "可承担各类核反应堆、放射性化工、核燃料元件、核同位素分离、铀冶金、核废料处理、核电站检修和维修以及铀矿山工程的施工。",
        "二级": "可承担合同额 6000 万以下的放射性化工、核燃料元件、核同位素分离、铀冶金、核废料处理、核电站检修和维修以及铀矿山工程的施工。"
    },
    "海洋石油工程专业承包": {
        "一级": "可承担各类型海洋石油工程和其他海洋工程的施工、维修、改造等。",
        "二级": "可承担项目投资额 8 亿元以下海洋油气开发工程或 3 亿元以下海底管道工程，以及其他海洋工程的施工、维修、改造等。"
    },
    "环保工程专业承包": {
        "一级": "可承担各类环保工程的施工。",
        "二级": "可承担污染修复工程、生活垃圾处理处置工程大型以下及其他中型以下环保工程的施工。",
        "三级": "可承担污染修复工程、生活垃圾处理处置工程中型以下及其他小型环保工程的施工。"
    }
}

COMPLIANCE_CHECKS = [
{
        "name": "限定营业执照经营范围",
        "description": "审查是否要求投标人的营业执照经营范围必须涵盖招标产品的制作或销售等特定范围。",
        "legal_basis": "《招标投标法实施条例》第三十二条：招标人不得以不合理的条件限制、排斥潜在投标人。设定的资格、技术、商务条件与招标项目的具体特点和实际需要不相适应或者与合同履行无关，属于不合理条件。"
    },
    {
        "name": "限定或指定特定专利、商标、品牌、原产地或供应商",
        "description": "审查是否限定或指定特定专利、商标、品牌、原产地或供应商，排斥其他潜在投标人。",
        "legal_basis": "《招标投标法实施条例》第三十二条：招标人不得以不合理的条件限制、排斥潜在投标人。"
    },
    {
        "name": "要求投标人注册资本金",
        "description": "审查是否在投标人资格要求中设置了注册资本金门槛（如'注册资本金在XXX万元以上'）。",
        "legal_basis": "《招标投标法实施条例》第三十二条：招标人不得以不合理的条件限制、排斥潜在投标人。设定的资格条件与招标项目的具体特点和实际需要不相适应，属于不合理条件。"
    },
    {
        "name": "度身定向招标或不合理条件限制、排斥潜在投标人",
        "description": "审查合同中是否存在度身定向招标、设定的资格/技术/商务条件与项目实际需要不相适应或与合同履行无关、排斥其他潜在投标人公平竞争的情形。",
        "legal_basis": "《招标投标法实施条例》第三十二条、第三十三条：招标人不得以不合理的条件限制、排斥潜在投标人；设定的资格、技术、商务条件与招标项目的具体特点和实际需要不相适应或者与合同履行无关，属于不合理条件；投标人参加依法必招项目不受地区或部门限制。"
    },


    {
        "name": "设立其他非法保证金",
        "description": "审查是否要求缴纳除投标保证金、履约保证金、工程质量保证金、农民工工资保证金以外的其他保证金。",
        "legal_basis": "《国务院办公厅关于清理规范工程建设领域保证金的通知》：除上述四种保证金外，其他一律取消，严禁新设保证金项目。"
    },
    {
        "name": "无息退还投标保证金或退还时间不合规",
        "description": "审查是否约定无息退还投标保证金，或未在合同签订后5日内退还。",
        "legal_basis": "《招标投标法实施条例》第五十七条：招标人最迟应当在书面合同签订后5日内向中标人和未中标人退还投标保证金及银行同期存款利息。"
    },
    {
        "name": "以不合理的理由没收投标保证金",
        "description": "审查是否规定了超出法定情形（如投标截止后撤销、中标后拒签合同等）的不合理没收投标保证金条件。",
        "legal_basis": "《招标投标法实施条例》：仅当投标人在投标截止后撤销投标文件，中标人无正当理由不与招标人订立合同，中标人提出附加条件或不按要求提交履约保证金时，招标人可以不退还投标保证金。"
    },

    {
        "name": "劳务分包招标文件中包含主材、建筑材料款、机械费、周转材料等",
        "description": "审查劳务分包招标是否包含主材、大型机械等费用，属于违法分包。",
        "legal_basis": "《建筑工程施工发包与承包违法行为认定查处管理办法》第十二条：专业作业承包人除计取劳务作业费用外，还计取主要建筑材料款和大中型施工机械设备、主要周转材料费用的，属于违法分包。"
    },
    {
        "name": "劳务分包招标文件原封不动引用主合同清单",
        "description": "审查劳务分包招标文件是否直接复制主合同的工程量清单，可能导致计价方式不合理。",
        "legal_basis": "合理商业惯例，避免将主合同风险转嫁给劳务分包。"
    },
    {
        "name": "分包招标文件编制中包含项目主体结构部分的清单",
        "description": "审查分包招标是否包含主体结构施工内容。",
        "legal_basis": "《中华人民共和国建筑法》第二十九条：施工总承包的，建筑工程主体结构的施工必须由总承包单位自行完成。"
    },
    {
        "name": "招标文件中规定单价与数量乘积与合价不一致时按就低不就高原则修正",
        "description": "审查是否设置了不合理的修正规则（如就低不就高）。",
        "legal_basis": "《评标委员会和评标办法暂行规定》第十九条：总价与单价不一致时以单价为准，大写与小写不一致时以大写为准。未规定就低不就高。"
    },
    {
        "name": "招标文件中设置开标谈判程序",
        "description": "审查是否允许在开标后与投标人进行谈判。",
        "legal_basis": "《招标投标法》第四十三条：在确定中标人前，招标人不得与投标人就投标价格、投标方案等实质性内容进行谈判。"
    },
    {
        "name": "开标后出现第二轮竞价或第二次谈判条款",
        "description": "审查是否允许二次报价或谈判。",
        "legal_basis": "《招标投标法》第四十三条：在确定中标人前，招标人不得与投标人就投标价格、投标方案等实质性内容进行谈判。且《工程建设项目施工招标投标办法》第五十九条禁止向中标人提出压低报价等要求。"
    },
    {
        "name": "无正当理由随意暂停或终止招标",
        "description": "审查是否允许招标人擅自暂停或终止招标。",
        "legal_basis": "《工程建设项目施工招标投标办法》第十五条：除不可抗力原因外，不得擅自终止招标。"
    },
    {
        "name": "要求不合理的签订合同时限",
        "description": "审查是否要求中标人在少于30日内签订合同。",
        "legal_basis": "《招标投标法》第四十六条：招标人和中标人应当自中标通知书发出之日起三十日内订立书面合同。"
    }
]

# ============================================
# 文本提取
# ============================================
def extract_text_from_docx(file_obj):
    doc = Document(file_obj)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)

def extract_text_fallback(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type == "pdf":
        reader = pypdf.PdfReader(uploaded_file)
        return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif file_type == "docx":
        return extract_text_from_docx(uploaded_file)
    else:
        return uploaded_file.read().decode("utf-8")

def extract_contract_text(uploaded_file):
    return extract_text_fallback(uploaded_file)

# ============================================
# 辅助函数
# ============================================
def normalize_text(s):
    if not isinstance(s, str):
        return ""
    s = re.sub(r'[^\w\u4e00-\u9fff\s]', '', s)
    s = re.sub(r'\s+', ' ', s).strip().lower()
    return s

def extract_numbers(s):
    match = re.search(r'(\d+(?:\.\d+)?)', s)
    return match.group(1) if match else None

def is_numeric_equal(num_str1, num_str2):
    if num_str1 is None or num_str2 is None:
        return False
    try:
        f1 = float(num_str1)
        f2 = float(num_str2)
        if f1 != f2:
            return False
        def has_nonzero_decimal(s):
            if '.' not in s:
                return False
            decimal_part = s.split('.')[1].rstrip('0')
            return len(decimal_part) > 0 and any(c != '0' for c in decimal_part)
        has1 = has_nonzero_decimal(num_str1)
        has2 = has_nonzero_decimal(num_str2)
        return has1 == has2
    except:
        return False

def ai_extract_in_area(text, field_name, field_description, area_name="", default="未找到"):
    zhipu_client = OpenAI(
        api_key=st.secrets["ZHIPU_API_KEY_1"],
        base_url="https://open.bigmodel.cn/api/paas/v4/"
    )

    if not text or len(text.strip()) < 10:
        return default
    
    # 构建提示词 - 让AI自己理解字段含义
    prompt=f"""你是一个专业的招标文件信息提取助手。请从以下{area_name}中提取指定的字段，要求准确理解语义，区分相似但不同的概念。

【目标字段】：{field_name}
【字段描述】：{field_description}

【提取规则】：
1. 根据"{field_description}"理解你要找的内容是什么
2. 文本中可能有措辞相近但意思不同的内容，请选择最匹配的
3. 比如"联系人"可能有多个，选描述中最匹配的那个
4. 如果文本中有同义词或近义词，也视为匹配（如"收件人"和"收件联系人"）
5. 字段可能以"："或空格与内容分隔
6. 字段名和内容之间可能有多个空格
7. 根据字段描述的含义来识别，而不是固定的词
8. 不要提取不相关的内容，只从提供的文本中提取，不要编造
9. 关于收件联系人的查找，必须同时提取"姓名"和"联系电话"，姓名和联系电话通常在同一段落或相邻为找。
10. 只提取字段的值，不要提取字段名本身
11. 如果找不到，输出"未找到"

【文本内容】：
{text[:8000]}

【输出】："""
    
    try:
        response = zhipu_client.chat.completions.create(
            model="glm-4-flash-250414",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=600,
            timeout=15,
            stream=False
        )
        
        if hasattr(response, 'usage'):
            usage = response.usage
            print(f"【Token用量】{field_name}:")
            print(f"  - 输入Token: {usage.prompt_tokens}")
            print(f"  - 输出Token: {usage.completion_tokens}")
            print(f"  - 总计Token: {usage.total_tokens}")

        result = response.choices[0].message.content.strip()
        if "未找到" in result or not result:
            return default
        return result
    except Exception as e:
        print(f"AI在{area_name}提取{field_name}失败: {e}")
        return default

def extract_qualification_advanced(text):
    """
    从资格要求文本中提取资质类别和等级。
    返回 (category, level) 元组，例如 ("建筑装修装饰工程专业承包", "二级")
    """
    if not text:
        return None, None
    
    # 1. 从字典中精确匹配类别
    all_categories = list(QUALIFICATION_STANDARDS.keys())
    # 按长度降序，优先匹配长名称
    all_categories_sorted = sorted(all_categories, key=len, reverse=True)
    found_category = None
    for cat in all_categories_sorted:
        if cat in text:
            found_category = cat
            break
    
    if not found_category:
        # 若字典中没有，尝试用正则提取
        cat_match = re.search(r'([\u4e00-\u9fff]+?(?:施工总承包|专业承包|劳务分包))', text)
        if cat_match:
            found_category = cat_match.group(1)
        else:
            return None, None
    
    # 2. 提取等级（支持中文数字和阿拉伯数字）
    level_map = {
        "一级": "一级", "壹级": "一级",
        "二级": "二级", "贰级": "二级",
        "三级": "三级", "叁级": "三级"
    }
    level = None
    for key, val in level_map.items():
        if key in text:
            level = val
            break
    
    return found_category, level

def check_with_ollama(category, level, desc, max_price, project_scale):
    if not category or not level:
        return False, "未提取到有效的资质类别或等级"
    if not desc:
        return False, f"未找到 {category} {level} 的官方描述"

    # 初始化客户端
    client = OpenAI(
        api_key=st.secrets["ZHIPU_API_KEY_2"],
        base_url="https://open.bigmodel.cn/api/paas/v4/"
    )

    prompt = f"""你是一位招标文件合规审查专家。请根据以下**已确定的资质类别和等级**的标准描述，判断招标文件中设定的该资质等级要求是否合理。

【已确定的资质】
- 类别：{category}
- 等级：{level}

【标准描述】（摘自官方资质标准，请以此为准）
{desc}

【本项目关键数据】
- 采购控制价（最高报价限价）：{max_price if max_price else "未提供"}
- 项目概况（从公告"采购项目概况"提取的原始描述）：{project_scale if project_scale else "未提供"}

**强制判断流程（严格按顺序执行）**：

#### 步骤一：识别限制条件
从标准描述中找出所有含数值的限制条件，并分类：
- **金额限制**：关键词包括"合同额"、"单项合同额"、"投资额"、"总价" 
- **规模限制**：关键词包括"建筑面积"、"高度"、"跨度"、"重量"等。

#### 步骤二：统一单位（只做单位换算，不进行比较）
对于每一个限制条件，提取其数值和单位（如"2000万元"、"15000平方米"），并从项目数据中提取对应的数值和单位（如采购控制价"1500000元"、项目概况中的"建筑面积30906.32㎡"）。

**单位换算规则**：
- **金额单位**：统一转换为"元"。
  - 如果单位是"万元"，将数值乘以 10000 变为"元"。
  - 如果单位是"亿元"，将数值乘以 100000000 变为"元"。
  - 清理千分位分隔符（如"1,500,000" → "1500000"）和空格。
- **面积单位**：统一转换为"平方米"。
  - 如果单位是"公顷"，乘以 10000；如果是"平方公里"，乘以 1000000。
  - 清理千分位分隔符和空格。
- **长度单位**：统一转换为"米"。
  - 如果单位是"公里"，乘以 1000。
- **重量单位**：统一转换为"千克"或"吨"，根据上下文选择（标准中通常为"吨"）。

**重要**：**在统一单位后，所有数值均为纯数字**。

#### 步骤三：限制处理（必须先匹配，再比较）

**⚠️ 核心警告：概念匹配时，禁止语义推断，必须逐字核对原文！**
**"隐含"不算出现！"建筑面积"不等于"单体建筑面积"！**

对于标准描述中每一个限制，**必须按以下顺序执行**：

**A. 金额限制处理**（如"单项合同额2000万元"）：
1. 检查项目数据中是否提供了"采购控制价"、"最高报价限价"、"最高限价"、"控制价"中的**任意一个**。
   - 若为"未提供" → 标记为"不适用"，禁止比较
   - 如果提供了**任意一个数值**（即使名称不完全相同），都视为**概念匹配成功**。
2. 单位统一后比较：采购控制价（最高报价限价） ≤ 限制数值 → 合规；> 限制数值 → 不合理

**B. 规模限制处理**（如"单体建筑面积15000平方米"）：
1. **逐字核对**：将限制关键词拆解为核心名词和前缀描述。
   - 例如限制关键词"单体建筑面积" → 前缀描述为单体，核心名词为建筑面积。
  

规模限制：

2. 对于每个规模限制，按以下格式创建证据清单：
    - 限制名称：[例如：单体建筑面积 15000平方米]
    - 核心名词：[例如：建筑面积]
    - 项目概况中是否出现核心名词？：[是/否] → [引用原文]
    - 前缀描述：[例如：单体]
    - 项目概况中是否出现前缀描述？：[是/否] → [引用原文]
    - 是否匹配：[是/否]（核心名词和前缀描述都出现并一致才为"是"）

3. **匹配规则**：
   - 限制关键词中的所有词语都必须在项目概况中出现 → 匹配成功
   - 任意一个词语不出现 → 匹配失败
   - **禁止**将"建筑面积"当成"单体建筑面积"（因为缺少"单体"）
   - **禁止**将"高度"当成"钢结构高度"（因为缺少"钢结构"）

**关键规则**：
- 金额限制和规模限制**互不影响**，各自独立判断。
- 采购控制价**只**用于金额限制比较，**不**用于规模限制比较。
- 项目概况数据**只**用于规模限制比较，**不**用于金额限制比较。
- 限制关键词中的每个词语，都必须在项目概况原文中**逐字出现**。
- **"隐含"、"推断"、"相当于"等都不算匹配**。


**步骤四：数值比较**
- 根据步骤三的第一步找到的“证据”中提取数值。
- 按步骤二的规则统一单位。

**比较规则**：
1. 取出项目数据数值（如 150 元）
2. 取出限制数值（如 2000 元）
3. 判断：项目数据 **≤** 限制数值 → 合规；项目数据 **>** 限制数值 → 不合理

**⚠️ 警告：以下操作仅对"匹配成功"的限制执行！**

- **如果某个限制在步骤三中被标记为"匹配失败"或"不适用"** → **立即跳过**，**禁止**进行任何数值比较
- **只有步骤三中标记为"匹配成功"的限制** → 才执行数值比较

**检查清单**（在比较前必须确认）：
1. [ ] 该限制在步骤三中标记为"匹配成功"？
2. [ ] 如果是 → 进行数值比较
3. [ ] 如果否 → 不比较，继续下一个限制




#### 步骤五：输出结论
请回答以下**所有问题**，然后基于回答输出最终 JSON。

---

**问题1：标准中有哪些限制条件？**
请列出标准描述中所有含数值的限制条件，并标明类型（金额/规模）。



**问题2：对于每个限制，项目数据中找到了什么？**
请按以下格式回答：

金额限制：
- 限制名称：[例如：单项合同额 2000万元]
- 项目数据对应字段：[采购控制价/最高报价限价]
- 提供的数值：[具体数值 或 "未提供"]
- 是否匹配：[是/否]（提供了数值即为"是"）

规模限制：
- 限制名称：[例如：单体建筑面积 15000平方米]
- 核心名词：[例如：建筑面积]
- 项目概况中是否出现核心名词？：[是/否] → [引用原文]
- 前缀描述：[例如：单体]
- 项目概况中是否出现前缀描述？：[是/否] → [引用原文]
- 是否匹配：[是/否]（核心名词和前缀描述都出现并一致才为"是"）

**问题3：匹配成功的限制，数值比较结果是什么？**
- [限制名称]：项目数据 [数值] [≤/>] 限制数值 [数值] → [合规/不合理]
- 如果没有匹配成功的限制，回答"无"

**问题4：最终结论是什么？**
- [所有限制均不适用  / 存在超出限制/所有匹配限制均合规]

---

**基于以上回答，输出最终 JSON**：


**三种情况**：
1. **所有限制均为"匹配失败"或"不适用"** → `compliant`: true，`reason`: "所有限制均不适用"

2. **有"匹配成功"且全部合规** → `compliant`: true，`reason` 必须列出所有匹配成功且合规的限制：
   - 格式：`"所有匹配限制均合规。×××限制合规（项目数据 ××× ≤ 限制数值 ×××"`

3. **有"匹配成功"且存在超出** → `compliant`: false，`reason` 必须列出超出限制的具体内容：
   - 格式：`"存在超出限制。×××限制超出（项目数据 ××× > 限制数值 ×××）"`

**重要**：`reason` 必须用中文，且开头必须为上述三种之一。



"""
    try:
        stream = client.chat.completions.create(
            model="glm-4-flash-250414",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=2000,
            extra_body={"enable_thinking": False},
            stream=True
        )
        
        full_response = ""
        estimated_output_tokens = 0
        estimated_input_tokens = len(prompt) // 2
        
        for chunk in stream:
            delta = chunk.choices[0].delta
            if hasattr(delta, "content") and delta.content:
                full_response += delta.content
                estimated_output_tokens += len(delta.content) // 2
        
        print(f"【Token用量】check_with_ollama:")
        print(f"  - 估算输入Token: {estimated_input_tokens}")
        print(f"  - 估算输出Token: {estimated_output_tokens}")
        print(f"  - 估算总计Token: {estimated_input_tokens + estimated_output_tokens}")



        import re
        json_match = re.search(r'\{[^{}]*\}', full_response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            return result.get("compliant", True), result.get("reason", "未提供理由")
        else:
            if "合理" in full_response:
                return True, "模型判断为合理"
            else:
                return False, f"模型判断为不合理，原始返回：{full_response[:200]}"
            
    except Exception as e:
        error_msg = str(e)
        print(f"\n❌ API调用错误：{error_msg}")
        return True, f"API调用失败，默认合理。错误：{error_msg}"
        
# 智能比较函数
# ============================================
def smart_compare(front_val, table_val, field_name):
    if front_val == "未找到" or table_val == "未找到":
        return False, front_val, table_val, "缺失数据"

    # 分供商/响应人资格要求：去除开头和结尾的序号及引导语后再比较
    if field_name in ["分供商资格要求", "响应人资格要求"]:
        def clean_qualification(text):
            # 1. 去除开头的"1、分供商应依法设立且满足如下要求："等
            text = re.sub(r'^[\d一二三四五六七八九十]+[、\.]\s*分供商应依法设立且满足如下要求[：:]\s*', '', text.strip())
            # 2. 去除开头的"（1）资质要求："等
            text = re.sub(r'^（\d+）[^：:]*[：:]\s*', '', text)
            # 3. 去除开头的纯序号（如"1、"）
            text = re.sub(r'^[\d一二三四五六七八九十]+[、\.]\s*', '', text)
            
            # 🔧 新增：去除行内的所有序号（如 11.1、11.2、11.3、（1）、① 等）
            # 去除 11.1、11.2 格式的序号
            text = re.sub(r'\s*\d+\.\d+\s*', ' ', text)
            # 去除 1.、2. 格式的序号
            text = re.sub(r'\s*\d+[\.、]\s*', ' ', text)
            # 去除（1）、（2）格式的序号
            text = re.sub(r'\s*（\d+）\s*', ' ', text)
            # 去除 (1)、(2) 格式的序号
            text = re.sub(r'\s*\(\d+\)\s*', ' ', text)
            # 去除 ①、② 格式的序号
            text = re.sub(r'\s*[①②③④⑤⑥⑦⑧⑨⑩]\s*', ' ', text)
            # 去除结尾的序号（含换行）
            text = re.sub(r'\s*[\d一二三四五六七八九十]+[、\.]\s*$', '', text)

                # 去除 （1）资格要求： 或 （1）承担本项目的主要人员要求： 等
            text = re.sub(r'（\d+）[^：:]*[：:]\s*', '', text)
    # 去除 (1)资格要求： 或 (1)承担本项目的主要人员要求： 等
            text = re.sub(r'\(\d+\)[^：:]*[：:]\s*', '', text)
    # 去除 1、资格要求： 或 1、承担本项目的主要人员要求： 等
            text = re.sub(r'[\d一二三四五六七八九十]+[、\.][^：:]*[：:]\s*', '', text)
    # 去除 "资格要求："、"人员要求："、"业绩要求："、"信誉要求："、"其他要求：" 等（不带序号）
            text = re.sub(r'[^：:]*[：:]\s*', '', text)
            
            # 5. 将多个空格、换行、制表符合并为单个空格
            text = re.sub(r'\s+', ' ', text)
            # 6. 清理中文标点后的多余空格（如"： " → "："）
            text = re.sub(r'[：:]\s+', '：', text)
            text = re.sub(r'；\s+', '；', text)
            text = re.sub(r'，\s+', '，', text)
            # 7. 去除括号内的多余空格（如"（ 5 ）" → "（5）"）
            text = re.sub(r'（\s*(\d+)\s*）', r'（\1）', text)
            # 8. 去除首尾空白
            return text.strip()
        
        front_clean = clean_qualification(front_val)
        table_clean = clean_qualification(table_val)
        
        if front_clean == table_clean:
            return True, front_clean, table_clean, "忽略序号和空格后一致"
        else:
            return False, front_clean, table_clean, "内容不一致"
    
    exact_fields = ["采购项目名称"]
    if field_name in exact_fields:
        front_clean = front_val.strip()
        table_clean = table_val.strip()
        if front_clean == table_clean:
            return True, front_clean, table_clean, "精确匹配"
        else:
            return False, front_clean, table_clean, "内容不一致"

    if field_name == "提交地点":
        def extract_address(text):  
            """提取纯地址，移除收件人、电话等附加信息"""
            # 先去掉收件人信息
            text = re.sub(r'[；;]\s*收件人[：:]\s*[^\s]+', '', text)
            text = re.sub(r'[；;]\s*电话[：:]\s*[\d\s-]+', '', text)
            text = re.sub(r'\s*收件人[：:]\s*[^\s]+', '', text)
            text = re.sub(r'\s*电话[：:]\s*[\d\s-]+', '', text)
            # 清理多余空格
            text = re.sub(r'\s+', ' ', text).strip()
            return text
        
        front_clean = extract_address(front_val)
        table_clean = extract_address(table_val)
        
        # 比较清理后的地址
        front_no_space = re.sub(r'\s+', '', front_clean)
        table_no_space = re.sub(r'\s+', '', table_clean)
        
        if front_no_space == table_no_space:
            return True, front_clean, table_clean, "忽略附加信息和空格后一致"
        else:
            norm_front = normalize_text(front_clean)
            norm_table = normalize_text(table_clean)
            if norm_front == norm_table or norm_front in norm_table or norm_table in norm_front:
                return True, norm_front, norm_table, "归一化或包含匹配"
            return False, front_val, table_val, "内容不一致"



    if field_name == "收件联系人":
        front_phone = re.search(r'\d{11}', front_val)
        table_phone = re.search(r'\d{11}', table_val)
        front_name = re.search(r'[\u4e00-\u9fff]{2,4}', front_val)
        table_name = re.search(r'[\u4e00-\u9fff]{2,4}', table_val)
        phone_ok = (front_phone and table_phone and front_phone.group(0) == table_phone.group(0))
        name_ok = (front_name and table_name and front_name.group(0) == table_name.group(0))
        if phone_ok and name_ok:
            return True, f"{front_name.group(0)} {front_phone.group(0)}", f"{table_name.group(0)} {table_phone.group(0)}", "姓名+手机号匹配"
        else:
            return False, front_val, table_val, "姓名或手机号不一致"


    if field_name == "建设规模":
        front_num_str = extract_numbers(front_val)
        table_num_str = extract_numbers(table_val)
        def normalize_unit(unit):
            if not unit:
                return ""
            unit_lower = unit.lower()
            if unit_lower in ['平方米', '㎡', 'm2', 'm2', '平方']:
                return "平方米"
            return unit_lower
        def extract_unit(s):
            match = re.search(r'(\d+(?:\.\d+)?)\s*(平方米|㎡|m2|m2|平方|m)', s)
            if match:
                return match.group(2)
            match = re.search(r'\d+(?:\.\d+)?\s*([^\d\s]{1,3})', s)
            if match:
                return match.group(1)
            return ""
        front_unit_raw = extract_unit(front_val)
        table_unit_raw = extract_unit(table_val)
        front_unit = normalize_unit(front_unit_raw)
        table_unit = normalize_unit(table_unit_raw)
        if front_num_str and table_num_str:
            if is_numeric_equal(front_num_str, table_num_str):
                if front_unit == table_unit:
                    return True, front_val, table_val, "数值+单位一致"
                else:
                    return False, front_val, table_val, f"数值相同但单位不同 (前部单位: {front_unit_raw}, 须知单位: {table_unit_raw})"
            else:
                return False, front_val, table_val, f"数值不一致 (前部: {front_num_str}, 须知: {table_num_str})"
        else:
            return False, front_val, table_val, "无法提取数值"
    
    if field_name in ["计划工期", "工期要求", "工期"]:
        def standardize_date(s):
            """统一日期格式为 YYYY年MM月DD日"""
            # 处理：2026年6月1日 → 2026年06月01日
            s = re.sub(
                r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
                lambda m: f"{m.group(1)}年{int(m.group(2)):02d}月{int(m.group(3)):02d}日",
                s
            )
            # 处理：2026-06-01 → 2026年06月01日
            s = re.sub(
                r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
                lambda m: f"{m.group(1)}年{int(m.group(2)):02d}月{int(m.group(3)):02d}日",
                s
            )
            return s
        
        # 先标准化日期
        front_std = standardize_date(front_val)
        table_std = standardize_date(table_val)
        
        # 传入的值已经是统一格式：合同工期总日历天数：XX天；开工日期：XXXX年XX月XX日；竣工日期：XXXX年XX月XX日
        front_clean = re.sub(r'\s+', '', front_std).strip()
        table_clean = re.sub(r'\s+', '', table_std).strip()
        
        # 完全一致
        if front_clean == table_clean:
            return True, front_clean, table_clean, "完全一致"
        
        # 提取天数、开工日期、竣工日期
        front_days = re.search(r'合同工期总日历天数[：:]\s*(\d+)\s*天', front_clean)
        table_days = re.search(r'合同工期总日历天数[：:]\s*(\d+)\s*天', table_clean)
        days_ok = (front_days and table_days and front_days.group(1) == table_days.group(1))
        
        front_start = re.search(r'开工日期[：:]\s*([^；;]+)', front_clean)
        table_start = re.search(r'开工日期[：:]\s*([^；;]+)', table_clean)
        start_ok = (front_start and table_start and front_start.group(1).strip() == table_start.group(1).strip())
        
        front_end = re.search(r'竣工日期[：:]\s*([^；;]+)', front_clean)
        table_end = re.search(r'竣工日期[：:]\s*([^；;]+)', table_clean)
        
        front_has_end = front_end and front_end.group(1).strip() != "未提供"
        table_has_end = table_end and table_end.group(1).strip() != "未提供"
        
        # 只有双方都有竣工日期时才比较
        if front_has_end and table_has_end:
            end_ok = (front_end.group(1).strip() == table_end.group(1).strip())
        else:
            end_ok = True  # 一方或双方没有竣工日期，视为通过
        
        # 判定：必须同时满足天数和开工日期一致
        if days_ok and start_ok:
            if front_has_end and table_has_end:
                if end_ok:
                    return True, front_clean, table_clean, "天数和开竣工日期均一致"
                else:
                    return False, front_clean, table_clean, "竣工日期不一致；天数和开工日期一致"
            else:
                return True, front_clean, table_clean, "天数和开工日期均一致"
        else:
            reasons = []
            if not days_ok:
                reasons.append("天数不一致")
            if not start_ok:
                reasons.append("开工日期不一致")
            if front_has_end and table_has_end and not end_ok:
                reasons.append("竣工日期不一致")
            return False, front_clean, table_clean, "；".join(reasons)

    if field_name == "报价及单价总价计价方式":
        keywords = ["固定综合单价", "固定单价", "综合单价", "可调价", "成本加酬金"]
        front_key = next((kw for kw in keywords if kw in front_val), None)
        table_key = next((kw for kw in keywords if kw in table_val), None)
        if front_key and table_key and front_key == table_key:
            return True, front_key, table_key, "关键词匹配"
        norm_front = normalize_text(front_val)
        norm_table = normalize_text(table_val)
        if norm_front in norm_table or norm_table in norm_front:
            return True, norm_front, norm_table, "归一化包含匹配"
        return False, front_val, table_val, "计价方式不一致"


    if field_name == "承包方式":
        contract_keywords = ["专业分包", "劳务分包", "材料采购", "服务分包", "机械租赁", "设备采购"]
        
        def extract_contract_type(text):
            for keyword in contract_keywords:
                if keyword in text:
                    return keyword
            return None
        
        front_type = extract_contract_type(front_val)
        table_type = extract_contract_type(table_val)
        
        if front_type and table_type and front_type == table_type:
            return True, front_type, table_type, f"承包方式一致（{front_type}）"
        else:
            return False, front_val, table_val, f"承包方式不一致（公告：{front_val}，须知：{table_val}）"
    if field_name in ["质量要求", "质量标准"]:
        front_clean = re.sub(r'\s+', ' ', front_val).strip()
        table_clean = re.sub(r'\s+', ' ', table_val).strip()
        
        if front_clean == table_clean:
            return True, front_clean, table_clean, "完全一致"
        if front_clean in table_clean or table_clean in front_clean:
            return True, front_val, table_val, "内容包含匹配"
        return False, front_clean, table_clean, "内容不匹配"



    if field_name == "截止时间":
        # 提取日期时间（允许数字与年月日时分之间有空格）
        pattern = r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*(\d{1,2})\s*时\s*(\d{1,2})\s*分'
        front_match = re.search(pattern, front_val)
        table_match = re.search(pattern, table_val)
        if front_match and table_match:
            front_clean = ''.join(front_match.groups())
            table_clean = ''.join(table_match.groups())
            if front_clean == table_clean:
                return True, front_clean, table_clean, "忽略空格匹配"
            else:
                return False, front_val, table_val, "内容不一致"
        else:
            # 回退：去除所有空格后比较整个字符串
            front_clean = re.sub(r'\s+', '', front_val)
            table_clean = re.sub(r'\s+', '', table_val)
            if front_clean == table_clean:
                return True, front_clean, table_clean, "忽略空格匹配"
            else:
                return False, front_val, table_val, "内容不一致"

    if field_name == "采购范围":
        import difflib
        def clean_scope(s):
            s = re.sub(r'（[^）]*）', '', s)
            s = re.sub(r'\([^)]*\)', '', s)
            s = re.sub(r'以[^。]*为准', '', s)
            s = re.sub(r'根据[^，,。]*[,，]?', '', s)
            s = re.sub(r'详见[^，,。]*[,，]?', '', s)
            s = re.sub(r'具体[^，,。]*[,，]?', '', s)
            s = re.sub(r'包含但不限于', '', s)
            s = re.sub(r'[，,。.、；;:：\s]+', '', s)
            s = re.sub(r'[^a-zA-Z\u4e00-\u9fff]', '', s)
            return s.strip()
        front_clean = clean_scope(front_val)
        table_clean = clean_scope(table_val)
        if not front_clean or not table_clean:
            return False, front_val, table_val, "内容为空"
        if front_clean in table_clean or table_clean in front_clean:
            return True, front_val, table_val, "内容包含匹配"
        ratio = difflib.SequenceMatcher(None, front_clean, table_clean).ratio()
        if ratio >= 0.6:
            return True, front_val, table_val, f"相似度匹配 ({ratio:.2f})"
        else:
            return False, front_val, table_val, f"内容不匹配 (相似度 {ratio:.2f})"

    norm_front = normalize_text(front_val)
    norm_table = normalize_text(table_val)
    if norm_front == norm_table:
        return True, norm_front, norm_table, "归一化匹配"
    front_num = extract_numbers(front_val)
    table_num = extract_numbers(table_val)
    if front_num and table_num and is_numeric_equal(front_num, table_num):
        return True, front_num, table_num, "数值匹配"
    return False, front_val, table_val, "不匹配"



# ============================================
# 解析前部信息（公告）
# ============================================
def parse_front_info_ai_only(text):
    """
    纯AI提取公告信息，跳过所有正则匹配
    """
    # 截取公告部分
    start_marker = "询价公告"
    end_marker = "目录"
    start_pos = text.find(start_marker)
    end_pos = text.find(end_marker, start_pos) if start_pos != -1 else -1
    if start_pos != -1:
        if end_pos != -1:
            text = text[start_pos:end_pos]
        else:
            text = text[start_pos:start_pos + 6000]
    else:
        text = text[:6000]
    
    info = {}
    match = re.search(r'采购项目名称[：:]\s*([^\n]+)', text)
    if match:
        info["采购项目名称"] = match.group(1).strip()
    else:
        match = re.search(r'工程名称[：:]\s*([^\n]+)', text)
        info["采购项目名称"] = match.group(1).strip() if match else "未找到"

    #降级，如果还找不到就用ai找
    if info.get("采购项目名称") in ["未找到", ""]:
        info["采购项目名称"] = ai_extract_in_area(
            text,
            "采购项目名称",
            "招标项目的名称",
            "公告"
        )    
    def ai_extract(prompt, field_name, default="未找到"):
        try:
            # 调用千问
            qianwen_client = OpenAI(
                api_key=st.secrets["ZHIPU_API_KEY_3"],
                base_url="https://open.bigmodel.cn/api/paas/v4/",
            )
            
            response = qianwen_client.chat.completions.create(
                model="glm-4-flash-250414",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=5000,
                extra_body={"enable_thinking": False},
                stream=False
            )
        
            result = response.choices[0].message.content.strip()
            if result and result != "未找到":
                return result
            return default
        except Exception as e:
            print(f"AI提取 {field_name} 失败: {e}")
            return default



    # ============================================
    # 2. 采购方式
    # ============================================
    prompt = f"""从以下公告文本中提取"采购方式"。

【规则】：
1. 查找"采购方式"、"招标方式"等标签
2. 也可以从公告标题或开头判断
3. 必须标准化为以下名称之一（只能输出这些）：
   - "公开询价"（匹配：公开邀请询价采购、询价采购）
   - "竞争性谈判（公开）"（匹配：竞争性谈判、公开竞争性谈判）
   - "邀请招标"（匹配：邀请招标）
   - "单一来源"（匹配：单一来源）
   - "公开招标"（匹配：公开招标）
4. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出标准化名称，不要有其他文字。"""
    info["采购方式"] = ai_extract(prompt, "采购方式")

    # ============================================
    # 3. 发文时间
    # ============================================
    prompt = f"""从以下公告文本中提取"发文时间"。

【规则】：
1. 优先查找"采购文件发布时间："标签后面的日期
2. 如果找不到，查找"公告发布时间："标签后面的日期
3. 日期格式为"XXXX年XX月XX日"（如"2026年6月16日"）
4. 注意：这不是"报名时间"，也不是"响应文件递交截止时间"
5. 只提取单个日期，不要提取时间段
6. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出日期（如"2026年6月16日"），不要有其他文字。"""
    info["发文时间"] = ai_extract(prompt, "发文时间")

    # ============================================
    # 4. 项目地点
    # ============================================
    prompt = f"""从以下公告文本中提取"项目地点"。

【规则】：
1. 优先查找"项目地点："标签后面的地址
2. 如果找不到，查找"建设地点："标签后面的地址
3. 地址通常包含省/市/区/路/街/巷/号等完整信息
4. 只提取地址本身，不要包含标签
5. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出完整地址（如"广东省广州市黄埔区绿地中央广场A3栋406"），不要有其他文字。"""
    info["项目地点"] = ai_extract(prompt, "项目地点")

    # ============================================
    # 5. 建设规模
    # ============================================
    prompt = f"""从以下公告文本中提取"建设规模"。



【你要找什么】：
建设规模就是这个项目"有多大"——比如建筑面积多少平方米、占地多少亩、工程体量多大。

【常见的表述方式】：
- 直接标签："建设规模："、"项目规模："、"工程规模："
- 在"采购项目概况"段落中描述
- 文中描述："总建筑面积约XX㎡"、"项目总用地面积XX公顷"

【提取原则】：
1. 优先找"建筑面积"或"总建筑面积"后面的数字+单位
2. 如果找不到建筑面积，找"建设规模"、"项目规模"、"工程规模"后面的数字+单位
3. 数字+单位是关键（如"30906.32㎡"、"约5万平方米"、"占地100亩"）
4. 如果看到"用地面积"但不是建筑面积，优先输出建筑面积
5. 如果找不到建筑面积，去找"采购项目概况"段落。提取该段标的内容
5. 如果实在找不到，输出"未找到"

【判断标准】：
- "总建筑面积30906.32㎡" → ✅ 这是建设规模
- "占地面积约50亩" → ⚠️ 这是用地面积，没有建筑面积时才用
- "项目总投资5000万元" → ❌ 这是投资额，不是建设规模
- "合同金额800万" → ❌ 这是金额，不是建设规模

【文本】：
{text}

【输出】：只输出建设规模（数字+单位），不要有其他文字。
输出示例1："30906.32㎡"
输出示例2："总建筑面积约5万平方米"
输出示例3："占地100亩"（只有在找不到建筑面积时才用"""

    info["建设规模"] = ai_extract(prompt, "建设规模")


    # ============================================
    # 6. 报价及单价总价计价方式
    # ============================================
    prompt = f"""从以下公告文本中提取"报价及单价总价计价方式"。

【你要找什么】：
报价方式就是"这个项目怎么算钱"——是固定总价？还是按单价乘以工程量？还是成本加酬金？

【常见的表述方式】：
- 直接标签："报价方式："、"承包方式："、"计价方式："
- 段落标题："报价方式"、"承包方式"（单独一行作为标题）
- 文中描述："本项目采用...方式报价"、"报价方式为..."

【提取原则】：
1. 找到描述"怎么算钱/怎么报价"的那段话
2. 不要纠结于有没有冒号、有没有换行，理解意思就行
3. 如果找到的内容很长（包含多个条款），只取最核心的那一句
4. 如果实在找不到，输出"未找到"

【判断标准】：
- 如果文本中说"固定总价"、"总价包干" → 这就是报价方式
- 如果文本中说"固定单价"、"综合单价"、"单价包干" → 这就是报价方式
- 如果文本中说"按实结算"、"成本加酬金" → 这也是报价方式
- 如果文本中只在说"怎么付款"、"怎么结算"（如"按月支付"），那不是报价方式

【文本】：
{text}

【输出】：只输出报价方式的简要描述（15-50字），不要有其他文字。
输出示例1："固定综合单价包干"
输出示例2："采用固定总价方式报价" 
输出示例3："固定单价，工程量按实结算"""

    info["报价及单价总价计价方式"] = ai_extract(prompt, "报价及单价总价计价方式")


#7.质量要求
    match = re.search(r'(?:工程)?(?:质量[、，]技术|质量要求|质量标准)[：:]\s*(.*?)(?=\n[^\n]*[：:]|\Z)', text, re.DOTALL)
    if match:
        info["质量要求"] = match.group(1).strip().replace('\n', ' ')[:500]
    else:
        prompt = f"""从以下公告文本中提取"质量要求"。

【你要找什么】：
质量要求就是这个项目"要做到什么标准"——比如合格、优良、符合什么规范、达到什么验收标准。

【常见表述方式】：
- "质量要求："、"质量标准："、"工程质量："
- "质量[、，]技术"这种变体
- 文中描述："工程质量符合...标准"

【提取原则】：
1. 找到描述质量要求的那段话
2. 如果找不到明确标签，就找包含"质量"+"合格/标准/规范/验收"等关键词的句子
3. 完整提取，不要截断
4. 如果实在找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出质量要求描述，不要有其他文字。"""
        info["质量要求"] = ai_extract(prompt, "质量要求")
    # ============================================
    # 8. 安全文明施工要求
    # ============================================
    prompt = f"""从以下公告文本中提取"安全文明施工要求"。

【规则】：
1. 查找"安全文明施工要求："标签后面的内容
2. 提取完整的安全生产和文明施工要求
3. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出安全文明施工要求，不要有其他文字。"""
    info["安全文明施工要求"] = ai_extract(prompt, "安全文明施工要求")

    # ============================================
    # 9. 采购范围
    # ============================================
    prompt = f"""从以下公告文本中提取"采购范围"

【规则】：
1. 查找"采购范围："标签后面的内容
2. 提取到下一个数字序号（如"1."、"一、"）或空行之前
3. 提取完整的采购内容和范围描述
4. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出采购范围描述，不要有其他文字。"""
    info["采购范围"] = ai_extract(prompt, "采购范围")

    # ============================================
    # 10. 计划工期
    # ============================================
    prompt = f"""从以下公告文本中提取"计划工期"。

【规则】：
1. 查找"计划工期："或"工期："标签后面的内容
2. 提取到下一个数字序号（如"10、"、"一、"）之前停止（跳过带括号的序号如（1））
3. 必须包含以下三项信息（如果存在）：
   - 开工日期：格式如"2026年6月1日"或"2026-06-01"
   - 竣工日期：格式如"2026年8月30日"或"2026-08-30"
   - 工期天数：如"60天"或"60日历天"
4. 输出格式：计划工期XX天；开工日期XXXX年XX月XX日；竣工日期XXXX年XX月XX日
5. 如果某项未找到，写"未提供"
6. 如果找不到任何工期信息，输出"未找到"

【文本】：
{text}

【输出示例】：计划工期60天；开工日期2026年6月1日；竣工日期2026年8月30日
【输出】：只输出工期信息，不要有其他文字。"""
    info["计划工期"] = ai_extract(prompt, "计划工期")

    # ============================================
    # 11. 报名时间
    # ============================================
    prompt = f"""从以下公告文本中提取"报名时间"。

【规则】：
1. 查找"报名时间："标签后面的内容
2. 提取完整的报名起止时间
3. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出报名时间，不要有其他文字。"""
    info["报名时间"] = ai_extract(prompt, "报名时间")

    # ============================================
    # 12. 分供商资格要求
    # ============================================
    prompt = f"""从以下公告文本中提取"分供商资格要求"。

【规则】：
1. 查找"分供商资格要求"、"响应人资格要求"、"供应商资格要求"等章节
2. **关键**：提取到**"分供商不得存在下列情形之一"**之前**立即停止**，**不要**提取该句及其后面的任何内容
3. 必须完整提取所有条目，不要遗漏，特别注意以下条目：
   - （1）资质要求：企业资质等级、安全生产许可证等
   - （2）业绩要求：类似项目业绩要求
   - （3）信誉要求：信用中国报告、无不良记录等
   - （4）承担本项目的主要人员要求：项目经理资格、建造师证书、安全B证、社保等
   - （5）其他要求：响应文件格式、签署要求等 **（⚠️ 特别注意：必须完整提取"其他要求"的全部内容，这是最容易遗漏的部分）**
4. 保留每个条目的标题（如"资质要求"、"人员要求"、"其他要求"等）
5. 如果原文只有一两句话，直接提取原文内容，不要编造不存在的条目
6. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出资格要求全文，不要有其他文字。"""
    info["分供商资格要求"] = ai_extract(prompt, "分供商资格要求")

    # ============================================
    # 13. 承包方式
    # ============================================
    prompt = f"""从以下公告文本中提取"承包方式"。

【规则】：
1. 从采购项目名称或公告内容中判断
2. 必须标准化为以下名称之一（只能输出这些）：
   - "专业分包"（包含"专业分包"、"专业"）
   - "劳务分包"（包含"劳务分包"、"劳务"）
   - "材料采购"（包含"材料采购"、"材料"）
   - "服务分包"（包含"服务分包"、"服务"）
   - "机械租赁"（包含"机械租赁"、"租赁"）
   - "设备采购"（包含"设备采购"、"设备"）
   - "施工分包"（包含"施工分包"）
3. 如果无法判断，输出"未找到"

【文本】：
{text}

【输出】：只输出标准化名称（如"专业分包"），不要有其他文字。"""
    info["承包方式"] = ai_extract(prompt, "承包方式")

    # ============================================
    # 14. 截止时间
    # ============================================
    prompt = f"""从以下公告文本中提取"截止时间"。

【规则】：
1. 优先在"响应文件的递交"章节中查找
2. 查找"响应文件递交的截止时间："或"截止时间："标签
3. 必须包含日期和具体时间（如"2026年7月7日15时00分"）
4. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出截止时间（如"2026年7月7日15时00分"），不要有其他文字。"""
    info["截止时间"] = ai_extract(prompt, "截止时间")

    # ============================================
    # 15. 提交地点
    # ============================================
    prompt = f"""从以下公告文本中提取"提交地点"。

【规则】：
1. 优先在"响应文件的递交"章节中查找
2. 查找"地点："、"提交地点："、"递交地点："等标签
3. 地址必须完整，包含省/市/区/路/街/巷/号等信息
4. 优先提取"提交至"或"递交至"后面的地址
5. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出完整地址（如"广东省广州市黄埔区绿地中央广场A3栋406"），不要有其他文字。"""
    info["提交地点"] = ai_extract(prompt, "提交地点")

    # ============================================
    # 16. 收件联系人
    # ============================================
    prompt = f"""从以下公告文本中提取"收件联系人"。

【规则】：
1. 在"联系方式"或"谈判活动联系方式"章节中查找
2. 优先查找"收件人："标签，其次"联系人："、"收件联系人："标签
3. 必须提取姓名和联系电话
4. 排除以下角色词（不是真正的联系人）：
   - "采购人"、"招标人"、"代理机构"、"业主"
   - "甲方"、"乙方"、"承包人"、"分包商"
   - "供应商"、"响应人"、"投标人"
5. 输出格式：姓名 电话
6. 如果找不到，输出"未找到"

【文本】：
{text}

【输出】：只输出"姓名 电话"格式，不要有其他文字。"""
    info["收件联系人"] = ai_extract(prompt, "收件联系人")

    # ============================================
    # 后处理：截断长字段
    # ============================================
    if info.get("计划工期") and info["计划工期"] != "未找到" and len(info["计划工期"]) > 50:
        info["计划工期"] = info["计划工期"][:50] + "..."

    if info.get("报价及单价总价计价方式") and info["报价及单价总价计价方式"] != "未找到" and len(info["报价及单价总价计价方式"]) > 50:
        info["报价及单价总价计价方式"] = info["报价及单价总价计价方式"][:50] + "..."

    # ============================================
    # 补充结构化工期字段（用于比较）
    # ============================================
    if info.get("计划工期") and info["计划工期"] != "未找到":
        plan_text = info["计划工期"]
        
        # 天数匹配
        days_match = re.search(r'(\d+)\s*(?:日历天|天)', plan_text)
        if days_match:
            info["计划工期_天数"] = f"合同工期总日历天数：{days_match.group(1)}天"
        else:
            info["计划工期_天数"] = "合同工期总日历天数：未提供"
        
        # 🔧 开工日期：支持有冒号和无冒号
        start_match = re.search(r'开工日期[：:]?\s*([^；;。，,\n]+)', plan_text)
        if not start_match:
            start_match = re.search(r'计划开工日期[：:]?\s*([^；;。，,\n]+)', plan_text)
        if start_match:
            info["计划工期_开工日期"] = f"开工日期：{start_match.group(1).strip()}"
        else:
            info["计划工期_开工日期"] = "开工日期：未提供"
        
        # 🔧 竣工日期：支持有冒号和无冒号
        end_match = re.search(r'竣工日期[：:]?\s*([^；;。，,\n]+)', plan_text)
        if not end_match:
            end_match = re.search(r'计划竣工日期[：:]?\s*([^；;。，,\n]+)', plan_text)
        if end_match:
            info["计划工期_竣工日期"] = f"竣工日期：{end_match.group(1).strip()}"
        else:
            info["计划工期_竣工日期"] = "竣工日期：未提供"
    return info
# ============================================
# 解析手动粘贴的表格
# ============================================
def parse_table_from_text(table_text):
    def clean_markdown(s):
        s = re.sub(r'\[\[(.*?)\]\](?:\.\{[^\}]+\})?', r'\1', s)
        s = re.sub(r'\[(.*?)\](?:\.\{[^\}]+\})?', r'\1', s)
        s = re.sub(r'\{[^\}]+\}', '', s)
        return s

    raw_text = table_text
    table_text = clean_markdown(table_text)
    lines = table_text.strip().split('\n')
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|'):
            if re.match(r'^[\|\s\-:]+$', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) >= 3:
                rows.append((cells[0], cells[2]))
        elif re.match(r'^\d+\s+', line):
            parts = line.split(None, 2)
            if len(parts) >= 3:
                rows.append((parts[1], parts[2]))
    
    table_dict = {}
    for k, v in rows:
        k_clean = re.sub(r'^\d+\.?\s*', '', k).strip()
        if k_clean:
            table_dict[k_clean] = v



    # 拆分复合字段“响应文件提交地点及截止时间”
    target_key = None
    for key in table_dict.keys():
        if "响应文件提交地点及截止时间" in key:
            target_key = key
            break
    if target_key:
        # 从原始文本中重新提取该字段的完整内容（直到下一个数字序号或空行）
        pattern = r'响应文件提交地点及截止时间\s*(.*?)(?=\n\s*\d+[\.\s]|\n\s*\n|$)'
        match = re.search(pattern, raw_text, re.DOTALL)
        if match:
            value = match.group(1).strip()
        else:
            value = table_dict[target_key]  # 降级
        extra = split_composite_field(value,raw_text)
        del table_dict[target_key]
        table_dict.update(extra)

    # 降级提取提交地点和截止时间（如果拆分不完整）
    if "提交地点" not in table_dict or table_dict["提交地点"] == "未找到":
        loc_match = re.search(r'提交地点[：:]\s*([^。\n]+)', raw_text)
        if loc_match:
            table_dict["提交地点"] = loc_match.group(1).strip()
    if "截止时间" not in table_dict or table_dict["截止时间"] == "未找到":
        time_match = re.search(r'截止时间[：:]\s*([^。\n]+)', raw_text)
        if time_match:
            table_dict["截止时间"] = time_match.group(1).strip()

    if "提交地点" not in table_dict or table_dict.get("提交地点") in ["未找到", ""]:
        ai_location = ai_extract_in_area(
            raw_text,
            "提交地点",
            "须知表中响应文件或投标文件的递交地址，通常包含省/市/区/路等完整地址信息",
            "须知表"
        )
        if ai_location != "未找到":
            table_dict["提交地点"] = ai_location
    
    if "截止时间" not in table_dict or table_dict.get("截止时间") in ["未找到", ""]:
        ai_deadline = ai_extract_in_area(
            raw_text,
            "截止时间",
            "须知表中响应文件递交的截止时间，通常包含日期和具体时间（如'2026年6月25日15时00分'）",
            "须知表"
        )
        if ai_deadline != "未找到":
            table_dict["截止时间"] = ai_deadline


    # 提取响应人资格要求（限定在须知表范围内）- 纯AI提取
    table_dict["响应人资格要求"] = ai_extract_in_area(
        raw_text,
        "响应人资格要求",
        "从须知表中提取响应人/投标人/供应商需要满足的全部资格条件。\n\n"
        "【提取规则】：\n"
        "1. 只从上方提供的【须知表文本】中提取，不要从其他地方补充或推断。\n"
        "2. 如果原文内容是'详见谈判公告'、'同公告要求'、'同上'、'详见公告'、'详见招标文件'等引用性描述，直接返回该引用内容。\n"
        "3. 如果原文有具体的资格条件，完整提取所有内容。\n"
        "4. ⚠️ **关键**：提取到**下一个数字序号标题**之前停止（如'12响应文件份数'、'13、报价最高限价'等）。\n"
        "5. 保留原文的原有格式，不要改变段落结构、标点或条目标题。\n"
        "6. 如果找不到，输出'未找到'。\n\n"
        
    )

    # 提取工期要求的完整内容
    target_key = None
    for key in table_dict.keys():
        if "工期要求" in key:
            target_key = key
            break
    if target_key:
        current_val = table_dict[target_key]
        if "合同工期" not in current_val:
            start_keyword = "工期要求"
            start_pos = raw_text.find(start_keyword)
            if start_pos != -1:
                sub_start = start_pos + len(start_keyword)
                if sub_start < len(raw_text) and raw_text[sub_start] in '：:':
                    sub_start += 1
                while sub_start < len(raw_text) and raw_text[sub_start] in ' \t':
                    sub_start += 1
                match_seq = re.search(r'\n\s*(\d+)[\s\t]', raw_text[sub_start:])
                if match_seq:
                    end_pos = sub_start + match_seq.start()
                else:
                    end_pos = len(raw_text)
                full_text = raw_text[sub_start:end_pos].strip()
                if full_text and len(full_text) > len(current_val):
                    table_dict[target_key] = full_text

    # 确保报价方式键存在
    if "报价以及单价和总价计算方式" not in table_dict:
        for k in table_dict.keys():
            if k == "报价方式":
                table_dict["报价以及单价和总价计算方式"] = table_dict[k]
                break
    if "报价以及单价和总价计算方式" not in table_dict:
        table_dict["报价以及单价和总价计算方式"] = ai_extract_in_area(
            raw_text,
            "报价方式",
            "从须知表中提取报价方式，判断是'单价包干'还是'总价包干'。"
            "关键词：'综合单价包干'、'固定综合单价'、'固定单价' → 单价包干；"
            "'固定总价'、'综合总价包干'、'总价包干' → 总价包干。"
            "只输出'单价包干'或'总价包干'。",
            "须知表"
        )


    if "建设规模" in table_dict:
        original = table_dict["建设规模"]
        # 尝试从原文中提取总建筑面积的数字（优先从 raw_text 中查找）
        match = re.search(r'总建筑面积[^0-9]*(\d+(?:\.\d+)?)\s*㎡', raw_text)
        if match:
            # 同时提取单位（㎡）
            table_dict["建设规模"] = match.group(1) + "㎡"
        else:
            # 如果 raw_text 中找不到，则从当前值中尝试提取数字（可能为用地面积）
            num_match = re.search(r'(\d+(?:\.\d+)?)\s*(㎡|平方米|m2|平米)', original, re.IGNORECASE)
            if num_match:
                # 保留原数字，但注明单位（假设是㎡）
              table_dict["建设规模"] = num_match.group(1)+ num_match.group(2)
            else:
                # 没有面积单位，直接返回原文
                table_dict["建设规模"] = original


    if "采购方式" in table_dict:
        method_text = table_dict["采购方式"]
    # 使用同样的映射标准化
        method_mapping = [
        (r'公开竞争性谈判|竞争性谈判（公开）|竞争性谈判\(公开\)|竞争性谈判', "竞争性谈判（公开）"),
        (r'邀请竞争性谈判|竞争性谈判（邀请）|竞争性谈判\(邀请\)', "竞争性谈判（邀请）"),
        (r'公开询价|询价采购', "公开询价"),
        (r'定向询价|指定询价', "定向询价"),
        (r'公开招标（依法必招）|公开招标（非依法必招）|公开招标', "公开招标"),
        (r'邀请招标', "邀请招标"),
        (r'单一来源|单一来源采购', "单一来源"),
        ]
        for pattern, standard_name in method_mapping:
            if re.search(pattern, method_text, re.IGNORECASE):
                table_dict["采购方式"] = standard_name
                break
        return table_dict

def split_composite_field(value,raw_text=""):
    result = {}
    lines = value.split('\n')
    for line in lines:
        line = line.strip()
        if '收件人' in line:
            # 匹配格式：收件人：王军 18312174610 或 收件人：王军 电话：18312174610
            match = re.search(r'收件人[：:]\s*([^\d]+?)\s*(\d{11})', line)
            if match:
                name = match.group(1).strip()
                phone = match.group(2)
                result['收件联系人'] = f"{name} {phone}"
            else:
                # 没有手机号时，仅取姓名
                match = re.search(r'收件人[：:]\s*(.+)', line)
                if match:
                    result['收件联系人'] = match.group(1).strip()
        elif '提交地点' in line:
            match = re.search(r'提交地点[：:]\s*(.+)', line)
            if match:
                result['提交地点'] = match.group(1).strip()
        elif '截止时间' in line:
            match = re.search(r'截止时间[：:]\s*(.+)', line)
            if match:
                result['截止时间'] = match.group(1).strip()
    # 保底提取（如果上述未找到）
    if '收件联系人' not in result:
        phone_match = re.search(r'(\d{11})', value)
        name_match = re.search(r'(?:收件人|联系人)[：:]\s*([\u4e00-\u9fff]{2,4})', value)
        if phone_match and name_match:
            result['收件联系人'] = f"{name_match.group(1)} {phone_match.group(1)}"
        elif phone_match:
            result['收件联系人'] = phone_match.group(1)
        elif name_match:
            result['收件联系人'] = name_match.group(1)

    if raw_text and ('收件联系人' not in result or '提交地点' not in result or '截止时间' not in result):
        # 收件联系人AI兜底
        if '收件联系人' not in result or result.get('收件联系人') in ["", "未找到"]:
            ai_contact = ai_extract_in_area(
                raw_text,
                "收件联系人",
                "负责接收响应文件的人员姓名和联系电话，格式为'姓名 电话'",
                "须知表"
            )
            if ai_contact != "未找到":
                result['收件联系人'] = ai_contact
        
        # 提交地点AI兜底
        if '提交地点' not in result or result.get('提交地点') in ["", "未找到"]:
            ai_location = ai_extract_in_area(
                raw_text,
                "提交地点",
                "响应文件的递交地址，包含省/市/区/路",
                "须知表"
            )
            if ai_location != "未找到":
                result['提交地点'] = ai_location
        
        # 截止时间AI兜底
        if '截止时间' not in result or result.get('截止时间') in ["", "未找到"]:
            ai_deadline = ai_extract_in_area(
                raw_text,
                "截止时间",
                "响应文件递交的截止时间，包含日期和具体时间",
                "须知表"
            )
            if ai_deadline != "未找到":
                result['截止时间'] = ai_deadline
    
    return result


# ============================================
# 比对映射（公告 vs 须知表）
# ============================================
def compare_with_rules(front_info, table_dict):
    mapping = [
        ("采购项目名称", "工程名称"),
        ("项目地点", "建设地点"),
        ("建设规模", "建设规模"),
        ("质量要求", "质量标准"),
        ("承包方式", "承包方式"),
        ("采购方式","采购方式"),
        ("采购范围", "采购范围"),
        ("计划工期", "工期要求"),
        ("分供商资格要求", "响应人资格要求"),
        ("报价及单价总价计价方式", "报价以及单价和总价计算方式"),
        ("提交地点", "提交地点"),
        ("截止时间", "截止时间"),
        ("收件联系人", "收件联系人")
    ]
    results = []
    
    for front_key, table_key in mapping:
        front_val = front_info.get(front_key, "")
        
        # 特殊处理：报价方式
        if front_key == "报价方式":
            table_val = table_dict.get("报价方式") or table_dict.get("报价以及单价和总价计算方式")
            if table_val is None:
                results.append({
                    "项目": front_key,
                    "文件前部内容": front_val,
                    "须知样表内容": "（未找到对应项）",
                    "状态": "⚠️ 缺失",
                    "显示值": f"公告: {front_val}\n须知: 未找到"
                })
                continue
        
        # 特殊处理：计划工期
        # 特殊处理：计划工期
        # 特殊处理：计划工期
        elif front_key == "计划工期":
            # 保存原始值用于显示
            front_raw = front_info.get("计划工期", "")
            
            # 公告：使用已提取的结构化字段组合用于比较
            front_days = front_info.get("计划工期_天数", "")
            front_start = front_info.get("计划工期_开工日期", "")
            front_end = front_info.get("计划工期_竣工日期", "")
            front_val = f"{front_days}；{front_start}；{front_end}"
            
            # 须知表：从完整字段中提取关键信息
            table_raw = None
            for tk, tv in table_dict.items():
                if tk == table_key or table_key in tk:
                    table_raw = tv
                    break
            
            if table_raw is None:
                results.append({
                    "项目": front_key,
                    "文件前部内容": front_raw,
                    "须知样表内容": "（未找到对应项）",
                    "状态": "⚠️ 缺失",
                    "显示值": f"公告: {front_raw}\n须知: 未找到"
                })
                continue
            
            # 先尝试正则提取
            table_days = "合同工期总日历天数：未提供"
            table_start = "开工日期：未提供"
            table_end = "竣工日期：未提供"
            
            # 提取天数（支持"天"和"日历天"）
            days_match = re.search(r'(\d+)\s*(?:日历天|天)', table_raw)
            if days_match:
                table_days = f"合同工期总日历天数：{days_match.group(1)}天"
            
            # 提取开工日期（支持多种格式）
            start_match = re.search(r'从\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', table_raw)
            if not start_match:
                start_match = re.search(r'工期\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', table_raw)
            if not start_match:
                start_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', table_raw)
            if start_match:
                y, m, d = start_match.groups()
                table_start = f"开工日期：{y}年{int(m):02d}月{int(d):02d}日"
            
            # 提取竣工日期（支持"至"开头）
            end_match = re.search(r'至\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', table_raw)
            if not end_match:
                end_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日(?!.*?至)', table_raw)
            if end_match:
                y, m, d = end_match.groups()
                table_end = f"竣工日期：{y}年{int(m):02d}月{int(d):02d}日"
            
            table_val = f"{table_days}；{table_start}；{table_end}"
            
            # 如果正则提取不到天数或开工日期，用AI兜底
            if days_match is None or start_match is None:
                ai_table = ai_extract_in_area(
                    table_raw,
                      "工期要求",
                "从须知表中的工期描述中提取关键信息，注意以下格式：\n"
                "1. 总日历天数：提取数字+天（如'30天'或'30日历天'）\n"
                "2. 开工日期：提取开始日期，可能以'从'、'暂定工期'、'开工日期'开头\n"
                "3. 竣工日期：提取结束日期，可能以'至'、'竣工日期'开头\n"
                "输出格式：合同工期总日历天数：XX天；开工日期：XXXX年XX月XX日；竣工日期：XXXX年XX月XX日\n"
                "示例：合同工期总日历天数：30天；开工日期：2026年06月15日；竣工日期：2026年07月16日\n"
                "如果某项未找到，写'未提供'。",
                "须知表"
                )
                if ai_table != "未找到":
                    ai_days_match = re.search(r'合同工期总日历天数[：:]\s*(\d+)\s*天', ai_table)
                    ai_start_match = re.search(r'开工日期[：:]\s*([^；;]+)', ai_table)
                    ai_end_match = re.search(r'竣工日期[：:]\s*([^；;]+)', ai_table)
                    
                    if ai_days_match and days_match is None:
                        table_days = f"合同工期总日历天数：{ai_days_match.group(1)}天"
                    if ai_start_match and start_match is None:
                        table_start = f"开工日期：{ai_start_match.group(1).strip()}"
                    if ai_end_match and end_match is None:
                        table_end = f"竣工日期：{ai_end_match.group(1).strip()}"
                    
                    table_val = f"{table_days}；{table_start}；{table_end}"
            
            # 设置显示值
            display_front = front_raw
            display_table = table_raw
        
        else:
            # 其他字段正常取
            table_val = None
            for tk, tv in table_dict.items():
                if tk == table_key or table_key in tk:
                    table_val = tv
                    break
            if table_val is None:
                results.append({
                    "项目": front_key,
                    "文件前部内容": front_val,
                    "须知样表内容": "（未找到对应项）",
                    "状态": "⚠️ 缺失",
                    "显示值": f"公告: {front_val}\n须知: 未找到"
                })
                continue
            
            display_front = front_val
            display_table = table_val
        
        is_match, norm_front, norm_table, reason = smart_compare(front_val, table_val, front_key)
        status = "✅ 一致" if is_match else "❌ 不一致"
        
        results.append({
            "项目": front_key,
            "文件前部内容": display_front,
            "须知样表内容": display_table,
            "状态": status,
            "显示值": f"公告: {display_front}\n须知: {display_table}\n比较依据: {reason}"
        })
    return results




def ai_rules_check(text, rules, is_secondary=False, snippet_text=""):
    if not text or len(text.strip()) < 10:
        return []
    
    # 压缩文本
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r' +', ' ', text)
    text = text.replace('\t', ' ')
    text = re.sub(r'^ +| +$', '', text, flags=re.MULTILINE)
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\s*\d+\s*$', line):
            continue
        if re.match(r'^[\s\-|:]+$', line):
            continue
        cleaned_lines.append(line)
    text = '\n'.join(cleaned_lines)
    
    if len(text) > 100000:
        text = text[:100000] + "\n...（文件已截断）"

    rules_text = ""
    for i, rule in enumerate(rules, 1):
        rules_text += f"""
╔═══════════════════════════════════════════════════════════════
║ 【规则 {i}】{rule['name']}
╠═══════════════════════════════════════════════════════════════
║ 【法规依据】{rule.get('legal_basis', '')}
║ 【审查要点】{rule.get('description', '')}
╠═══════════════════════════════════════════════════════════════
"""
    zhipu_client2 = OpenAI(
        api_key=st.secrets["ZHIPU_API_KEY_3"],
        base_url="https://open.bigmodel.cn/api/paas/v4/",
        timeout=60.0    
    )
    rules_names = [rule["name"] for rule in rules]
    rules_names_text = "\n".join([f"- {name}" for name in rules_names])
    


    prompt = f"""你是一位招标文件合规审查专家。请根据以下五步审查法，对招标文件进行逐条审查。



## 🎯 五步审查法

请严格按照以下五个步骤对每条规则进行审查：


### 第一步：理解立法目的（为什么这条法规存在？）

这是审查的**出发点**。

**你要做什么：**
先搞清楚这条法规要防止什么、保护什么。

**你需要回答的问题：**
- 立法者为什么要制定这条规则？
- 这条规则想防止什么不好的事情发生？
- 这条规则想保护谁的利益？

**为什么这一步重要：**
- 立法目的是判断的**根本依据**
- 如果条款违背了立法目的，即使措辞不同也属于违规
- 如果不理解立法目的，就容易机械比对文字，漏掉变相违规

**【示例】**
- 规则："开标后出现第二轮竞价或第二次谈判条款"
- 立法目的：保证投标的公平性，预防招标人对投标人提出压低报价、增加工作量、缩短工期等要求。"

### 第二步：理解判定原则（什么算违规？）

这是审查的**核心标准**。

**你要做什么：**
从法规中提炼出"什么情况算违规"的本质特征。

**你需要回答的问题：**
- 这条法规划定的红线在哪里？
- 什么样的行为、条款、数字、情形超出了法规允许的范围？
- 违规的本质特征是什么？

**为什么这一步重要：**
- 判定原则给出了违规的**本质特征**
- 只要符合判定原则，无论具体措辞如何，都算违规
- 判定原则是连接"立法目的"和"具体条款"的桥梁

**【示例】**
"无息退还投标保证金"的判定原则是：
- 退还时间超过5日 → 违规
- 未约定利息或约定无息 → 违规
- 起算点不是"合同签订之日" → 违规

"招标文件中设置开标谈判程序"的判定原则：
-"合理低价法"、"综合评分法"、"经评审的最低投标价法"是评审办法，不是谈判程序，不构成违规
- 如果文件中没有出现"沟通"、"协商"、"谈判"、"交流"等词语，或没有明确的开标后交流程序，不得仅凭推测判定违规
- 开标时间与投标截止时间在同一时间进行 → 这是法规要求的正常程序，**不构成违规**
- 开标后（含"响应文件开启后"），招标人与投标人就**技术方案、报价组成、施工方法、人员配置、工期等实质性内容**进行任何形式的交流（包括但不限于"沟通"、"协商"、"交流"、"澄清"、"询问"），均属于变相谈判
- 只有对**格式、签字、盖章等非实质性内容**的澄清，才不构成谈判


### 第三步：参考常见违规举例（违规通常长什么样？）

这是审查的**辅助参考**。

**你要做什么：**
了解违规条款通常长什么样，帮助你在文件中快速识别。

**你需要知道：**
- 举例是为了帮助识别，**不是穷举**
- 如果文件中出现了举例中的模式 → 直接判定违规
- 如果文件中出现了与举例不同、但同样违背立法目的和判定原则的条款 → 也判定违规

**为什么这一步重要：**
- 帮助快速识别典型违规
- 但不要被举例限制住，要看**实质效果**

**【示例】**
- 规则1："设立其他非法保证金"
- 常见违规："诚信担保"、"合作保证金"、"安全保证金"、"廉政保证金"等

- 规则2："要求不合理的签订合同时限"
- 常见违规：
  - "中标通知书发出后45日内签订" → 45 > 30，无正当理由拖延 → 违规
- 合规举例：
  - "中标通知书发出后30日内签订" → 完全符合法规 → 合规


###第四步：概念区分（避免误判漏判）

这是防止审查错误的机智

**常见概念混淆**

1. **"合理低价法"、"综合评分法"、"经评审的最低投标价法"** 是评审办法，不是谈判程序，不构成违规

2. **主体结构**指建筑的承重结构：地基、梁、柱、承重墙、楼板、屋顶等
   - 只有主体结构工程必须由总承包单位自行完成
   - 装修工程、机电安装、幕墙、临时设施等不属于主体结构，可以分包，不构成违规

3. **"区分由谁提供/承担"决定是否违规**（针对劳务分包）：
   - "主要建筑材料由承包人提供/承担" → 分包人没有计取材料费 → 不构成违规
   - "主要建筑材料由分包人自行采购" → 分包人计取了材料费 → 构成违规

4. **"沟通"、"协商"、"交流"、"澄清"在开标后/响应文件开启后都可能构成变相谈判**
   - 仅对非实质性内容（格式、签字等）进行澄清，不构成违规
   - 针对投标价格、技术方案、施工方法、人员配置、工期等实质性内容进行交流，均属于谈判
   - "技术细节沟通"、"技术方案交流"等表述，在开标后发生且涉及实质性内容，构成违规
   - 评审过程是在响应文件开启/开标之后进行的，因此，"评审过程中"的任何实质性交流，都属于"开标后"的谈判

5. **区分"投标人"和"中标人"**：
    - "投标人" → 中标前的潜在中标者
          - 不得与投标人就价格、方案进行谈判 → 违规
          - 文件写了"合理低价法"、"报价最低者中标" → 这是评审办法，不涉及谈判 → 合规
    - "中标人" → 已经确定的中标者
          - 不得向中标人提出压低报价等要求 → 违规
          - 文件只提到"报价最低"而未涉及"与投标人沟通"、"压低报价" → 合规

6. 以下术语在招标文件中含义相同：
    - "响应文件开启" = "开标"
    - "响应文件递交截止" = "投标截止"
    - "成交人" = "中标人"
    - "分供商" = "投标人/供应商"
    - "采购人" = "招标人/业主"
    - "承包人" = "总包单位"



### 第五步：做出判断

这是审查的**结论**。

**你要做什么：**
综合前四步的分析，做出是否违规的判断。

**判断标准：**
- 不要求违规条款必须与"常见违规举例"完全一致
- 只要条款的**实质效果**符合"判定原则"，就应判定为违规
- 如果文件中没有相关条款，判定为合规


【以下是待审查的文件内容】
{text}

【以下是审查规则】
{rules_text}
【审查规则结束】


【审查要求】
1. 必须在上述文件中逐句查找违规证据
2. 如果找到了违规条款，引用原文作为 evidence
3. 如果文件中没有相关条款，evidence 写"无"
4. 不要只描述规则本身，必须基于文件内容做出判断
## 📤 输出格式
*重要：只能输出以下规则名称，不能输出列表之外的任何规则名称。**

【基础术语映射（审查前必读）】




【可输出的规则名称列表】
{rules_names_text}

对每条规则，按五步审查法分析后，输出JSON数组：

[
    {{
        "rule": "必须从上方【可输出的规则名称列表】中完整复制，不要添加'规则名称：'等任何前缀或后缀"
        "violation": true/false, // true=违规（文件中有问题），false=合规（文件中没有问题）
        "evidence": "必须填写在文件中找到的原文；如果在文件中没有找到，写'无'；禁止只写'返回原文'等无效内容",
        "reason": "判断理由，必须根据实际审查结果写明：\n
                   - 如果合规（violation=false）：'文件中未发现【规则名称】的相关条款，因此合规'（必须用实际规则名称替换【规则名称】）\n
                   - 如果违规（violation=true）：'文件中出现了违规原文（引用具体内容），违反了判定原则（描述具体违规原因），因此违规'"
]

**重要判断规则（必须遵守）**：


1. **前提条件判断**：有些规则仅在特定承包方式下才适用
   - 如果规则描述中包含"劳务分包"、"专业分包"等承包方式限定词，必须**先确认该项目属于该承包方式**
   - 如果不属于该承包方式，直接判定为**合规（不适用）**，不需要继续检查
   - 例如："劳务分包招标文件原封不动引用主合同清单" → 先确认是劳务分包，再检查清单引用

2. **违规判定标准**：
   - 如果文件中**没有**发现相关条款 → violation=false（合规）
   - 如果文件中**有**相关条款但条款内容明确**禁止**违规行为 → violation=false（合规）
   - 如果文件中**有**相关条款且条款内容明确**允许**违规行为 → violation=true（违规）
   - 如果 evidence 是"无"，violation 必须为 false
   - 如果 evidence 有内容（违规原文），violation 必须为 true

3. **判断步骤**：
   - 第一步：在文件中查找承包方式（如"承包方式"、"分包类型"、"合同类型"等）
   - 第二步：如果规则有前提条件，判断承包方式是否匹配
   - 第三步：如果承包方式不匹配，直接跳过，不再继续检查
   - 第四步：只有承包方式匹配时，才检查具体违规内容

【输出格式要求】：
- 只输出JSON数组，不要输出任何思考过程、解释、分析或其他文字。
- JSON中的字符串值不要包含双引号，如果必须包含，请使用单引号。
- 不要在JSON中使用中文标点符号。
- 每个规则必须包含 rule、violation、evidence、reason 四个字段。
- 不要输出"Here's a thinking process"或类似内容。
"""


    try:

        stream = zhipu_client2.chat.completions.create(
            model="glm-4-flash-250414",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=16000,
            extra_body={"enable_thinking": True},  # 开启思考过程
            stream=True  # 流式输出
        )
        full_response = ""
        estimated_input_tokens = len(prompt) // 2
        estimated_output_tokens = 0
        is_reasoning = True  # 标记是否在思考阶段

        for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            
            
            # 读取最终回复（content）
            if hasattr(delta, "content") and delta.content:
                full_response += delta.content
                estimated_output_tokens += len(delta.content) // 2

        
        print(f"【Token用量】ai_rules_check (共{len(rules)}条规则):")
        print(f"  - 估算输入Token: {estimated_input_tokens}")
        print(f"  - 估算输出Token: {estimated_output_tokens}")
        print(f"  - 估算总计Token: {estimated_input_tokens + estimated_output_tokens}")
        
        # 解析JSON
        start = full_response.find('[')
        end = full_response.rfind(']')
        if start != -1 and end != -1:
            json_str = full_response[start:end+1]
            
            last_bracket = json_str.rfind(']')
            if last_bracket != -1:
                json_str = json_str[:last_bracket+1]
            json_str = json_str.strip()
            
            results = []
            pattern = r'\{[^{}]*"rule"\s*:\s*"[^"]*"[^{}]*\}'
            matches = re.findall(pattern, json_str)
            
            for match in matches:
                try:
                    match_clean = match.strip()
                    match_clean = re.sub(r',\s*}', '}', match_clean)
                    result = json.loads(match_clean)
                    
                    # 清理 rule 字段中的前缀
                    rule_name = result.get("rule", "")
                    rule_name = re.sub(r'^规则名称[：:]\s*', '', rule_name)
                    rule_name = re.sub(r'^规则[：:]\s*', '', rule_name)
                    result["rule"] = rule_name
                    
                    results.append(result)
                except json.JSONDecodeError:
                    continue
            
            if not results:
                return []
            
            # 去重
            seen_rules = set()
            unique_results = []
            for result in results:
                rule_name = result.get("rule", "")
                if rule_name and rule_name not in seen_rules:
                    seen_rules.add(rule_name)
                    unique_results.append(result)
            results = unique_results
            
            # 只保留请求的规则
            expected_rules = [r["name"] for r in rules]
            filtered_results = []
            for result in results:
                if result.get("rule", "") in expected_rules:
                    filtered_results.append(result)
            results = filtered_results
            
            # 补充缺失的规则
            returned_rules = [r.get("rule", "") for r in results]
            for expected in expected_rules:
                if expected not in returned_rules:
                    results.append({
                        "rule": expected,
                        "violation": False,
                        "evidence": "未找到",
                        "reason": "AI返回被截断，默认合规"
                    })
            
            # 后处理：修正误判 + 补充缺失的reason
            for result in results:
                rule_name = result.get("rule", "")
                evidence = result.get("evidence", "")
                reason = result.get("reason", "")
                violation = result.get("violation", False)
                
                # 补充缺失的 reason
                if not reason or "描述具体违规原因" in reason or "XXX" in reason:
                    if violation:
                        result["reason"] = f"文件中出现了违规内容，具体见evidence字段"
                    else:
                        result["reason"] = f"未发现违规，合规"
                    reason = result["reason"]
                
                # 修正：竞争性谈判是采购方式，不是谈判程序
                if rule_name in ["招标文件中设置开标谈判程序", "开标后出现第二轮竞价或第二次谈判条款"]:
                    if "竞争性谈判" in evidence and not re.search(r'谈判.*?程序|谈判.*?沟通|谈判.*?协商|第二轮竞价|二次报价', evidence):
                        result["violation"] = False
                        result["reason"] = '"竞争性谈判"是合法的采购方式名称，不是开标后的谈判程序，合规'
                        continue
                
                # 情况1：没有证据 → 合规
                if evidence == "无" or "未发现" in evidence:
                    result["violation"] = False
                    result["reason"] = "未发现违规，合规"
                    continue
                
                # 情况2：有证据，检查 violation 和 reason 是否一致
                if evidence and evidence != "无":
                    reason_is_compliant = "合规" in reason and "违规" not in reason
                    reason_is_violation = "违规" in reason
                    
                    if violation and reason_is_compliant:
                        result["violation"] = False
                        result["reason"] = reason + "（合规）"
                    elif not violation and reason_is_violation:
                        result["violation"] = True
                        result["reason"] = reason + "（违规）"
                    else:
                        result["violation"] = violation
                        result["reason"] = reason
                else:
                    result["violation"] = False
                    result["reason"] = reason + "（无证据，合规）"
            
            return results
        return []
    except Exception as e:
        print(f"合并审查失败: {e}")
        return []

    """
    try:
        # 流式调用，开启思考过程
        stream = zhipu_client2.chat.completions.create(
            model="glm-4-flash-250414",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=10000,
            extra_body={"enable_thinking": True},  # 开启思考过程
            stream=True  # 流式输出
        )
        
        print("\n" + "=" * 20 + " 思考过程 " + "=" * 20)
        full_response = ""
        is_answering = False

        for chunk in stream:
            delta = chunk.choices[0].delta
            if hasattr(delta, "reasoning_content") and delta.reasoning_content is not None:
                print(delta.reasoning_content, end="", flush=True)
            if hasattr(delta, "content") and delta.content:
                if not is_answering:
                    print("\n" + "=" * 20 + " 最终回答 " + "=" * 20)
                    is_answering = True
                print(delta.content, end="", flush=True)
                full_response += delta.content
        
        print("\n" + "=" * 40)
        
        # 解析JSON

        start = full_response.find('[')
        end = full_response.rfind(']')
        if start != -1 and end != -1:
            results = json.loads(full_response[start:end+1])
            
            # 一致性校验
            for result in results:
                violation = result.get("violation", False)
                evidence = result.get("evidence", "")
                reason = result.get("reason", "")
                
                # 如果 AI 判合规，但证据中明确出现违规关键词，且理由中说的是"合规" → 才修正
                if violation == False and evidence and evidence != "无":
                    # 违规关键词列表
                    violation_keywords = ["指定", "必须使用", "仅限于", "承包人指定", "强制", "诚信保证金", "安全风险基金", "无息", "谈判", "协商"]
                    if any(kw in evidence for kw in violation_keywords) and "合规" in reason:
                        result["violation"] = True
                        result["reason"] = reason.replace("合规", "违规") + "（已修正：证据显示违规）"
                    else:
                        # 保持 AI 判断
                        result["violation"] = False
                        result["reason"] = reason + "（合规）"
                elif violation == True and evidence and evidence != "无":
                    # AI 判违规，保持
                    result["violation"] = True
                    result["reason"] = reason + "（违规）"
                else:
                    # 没有证据 → 合规
                    result["violation"] = False
                    result["reason"] = reason + "（无证据，合规）"
            return results
        return []
    except Exception as e:
        print(f"合并审查失败: {e}")
        return []
    """



# ============================================
# 内容合理性检查
# ============================================
def check_internal_rules(front_info, table_dict):
    issues = []
    skip_table = (table_dict == {})

    # 规则：采购方式与报名时间
    procurement_method = table_dict.get("采购方式", "")
    if procurement_method in ["", "未找到"]:
        procurement_method = front_info.get("采购方式", "")
    register_time = front_info.get("报名时间", "")
    method_days = {
        "公开询价": 2,
        "竞争性谈判（公开）": 3,
        "邀请招标": 5,
        "公开招标（依法必招）": 7,
        "公开招标（非依法必招）": 5,
    }
    if procurement_method and procurement_method != "未找到":
        matched = None
        for method, need_days in method_days.items():
            if method in procurement_method:
                matched = need_days
                break
        if matched is not None:
            if register_time and register_time != "未找到":
                date_pattern = r'(\d{4}年\d{1,2}月\d{1,2}日)'
                dates = re.findall(date_pattern, register_time)
                if len(dates) == 2:
                    start_str, end_str = dates
                    try:
                        start = datetime.strptime(start_str, "%Y年%m月%d日")
                        end = datetime.strptime(end_str, "%Y年%m月%d日")
                        days = (end - start).days+1
                        if days >= matched:
                            issues.append({
                                "规则": "报名时间",
                                "通过": True,
                                "原文": f"采购方式：{procurement_method}；报名时间：{register_time}",
                                "详情": f"报名时长为 {days} 天，满足最低要求 {matched} 天"
                            })
                        else:
                            issues.append({
                                "规则": "报名时间",
                                "通过": False,
                                "原文": f"采购方式：{procurement_method}；报名时间：{register_time}",
                                "详情": f"报名时长仅 {days} 天，低于最低要求 {matched} 天"
                            })
                    except:
                        issues.append({
                            "规则": "报名时间",
                            "通过": False,
                            "原文": f"采购方式：{procurement_method}；报名时间：{register_time}",
                            "详情": "报名时间格式解析失败"
                        })
                else:
                    issues.append({
                        "规则": "报名时间",
                        "通过": False,
                        "原文": f"采购方式：{procurement_method}；报名时间：{register_time}",
                        "详情": "报名时间不是有效的起止日期区间"
                    })
            else:
                issues.append({
                    "规则": "报名时间",
                    "通过": False,
                    "原文": f"采购方式：{procurement_method}",
                    "详情": "未找到报名时间"
                })
    else:
        issues.append({
            "规则": "报名时间",
            "通过": False,
            "原文": "未找到采购方式",
            "详情": "未在须知表中找到采购方式"
        })
    
    # 规则：投标时间
    issue_date = front_info.get("发文时间", "")
    deadline_str = front_info.get("截止时间", "")
    procurement_method = table_dict.get("采购方式", "")
    if procurement_method in ["", "未找到"]:
        procurement_method = front_info.get("采购方式", "")
    if issue_date != "未找到" and deadline_str != "未找到" and procurement_method and procurement_method != "未找到":
        try:
            issue_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', issue_date)
            if not issue_match:
                raise ValueError(f"发文时间格式错误: {issue_date}")
            iy, im, iday = issue_match.groups()
            issue_clean = f"{iy}年{im}月{iday}日"
            deadline_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', deadline_str)
            if not deadline_match:
                raise ValueError(f"截止时间未找到日期部分: {deadline_str}")
            dy, dm, dday = deadline_match.groups()
            deadline_clean = f"{dy}年{dm}月{dday}日"
            start = datetime.strptime(issue_clean, "%Y年%m月%d日")
            end = datetime.strptime(deadline_clean, "%Y年%m月%d日")
            days = (end - start).days+1
            method_min_days = {
                "公开询价": 1,
                "竞争性谈判（公开）": 2,
                "邀请招标": 2,
                "定向询价": 3,
                "竞争性谈判（邀请）": 5,
                "公开招标（依法必招）": 13,
                "公开招标（非依法必招）": 2,
                "单一来源": 1,
            }
            need_days = None
            for method, days_needed in method_min_days.items():
                if method in procurement_method:
                    need_days = days_needed
                    break
            if need_days is not None:
                if days >= need_days:
                    issues.append({
                        "规则": "投标时间",
                        "通过": True,
                        "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}，采购方式：{procurement_method}",
                        "详情": f"投标时间 {days} 天，满足最低要求 {need_days} 天"
                    })
                else:
                    issues.append({
                        "规则": "投标时间",
                        "通过": False,
                        "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}，采购方式：{procurement_method}",
                        "详情": f"投标时间仅 {days} 天，低于最低要求 {need_days} 天"
                    })
            else:
                issues.append({
                    "规则": "投标时间",
                    "通过": False,
                    "原文": f"采购方式：{procurement_method}",
                    "详情": "未知的采购方式，无法判断"
                })
        except Exception as e:
            issues.append({
                "规则": "投标时间",
                "通过": False,
                "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}",
                "详情": f"日期解析失败：{str(e)}"
            })
    else:
        missing = []
        if issue_date == "未找到":
            missing.append("发文时间")
        if deadline_str == "未找到":
            missing.append("截止时间")
        if not procurement_method or procurement_method == "未找到":
            missing.append("采购方式")
        issues.append({
            "规则": "投标时间",
            "通过": False,
            "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}，采购方式：{procurement_method}",
            "详情": f"缺少必要字段：{', '.join(missing)}"
        })
    
    # 规则：截止时间格式完整性检查
    # 规则：截止时间格式完整性检查
    deadline_front = front_info.get("截止时间", "")
    deadline_table = table_dict.get("截止时间", "")
    errors = []
    pattern = r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日\s*\d{1,2}\s*时\s*\d{1,2}\s*分'
    
    # 检查公告中的截止时间（仅当存在时）
    if deadline_front and deadline_front != "未找到":
        if not re.search(pattern, deadline_front):
            errors.append("公告中的截止时间缺少年月日时分中的部分成分")
    
    # 检查须知表中的截止时间（仅当存在时）
    if deadline_table and deadline_table != "未找到":
        if not re.search(pattern, deadline_table):
            errors.append("须知表中的截止时间缺少年月日时分中的部分成分")
    
    if errors:
        issues.append({
            "规则": "截止时间格式完整性",
            "通过": False,
            "原文": f"公告: {deadline_front}\n须知: {deadline_table}",
            "详情": "；".join(errors)
        })
    else:
        # 如果两个都不存在或都完整，视为通过
        issues.append({
            "规则": "截止时间格式完整性",
            "通过": True,
            "原文": f"公告: {deadline_front}\n须知: {deadline_table}",
            "详情": "截止时间格式完整" if (deadline_front != "未找到" or deadline_table != "未找到") else "未找到截止时间，跳过检查"
        })


    # 2. 履约担保比例（仅在未跳过表格时检查）
    if not skip_table:
        guarantee_text = ""
        # 尝试多种可能的键名
        for key in ["履约担保", "保函类型","履约保函"]:
            if key in table_dict:
                guarantee_text = table_dict[key]
                break
        if guarantee_text and guarantee_text != "未找到":
            match = re.search(r'(\d+(?:\.\d+)?)\s*%', guarantee_text)
            if match:
                percent = float(match.group(1))
                if percent >= 10:
                    issues.append({
                        "规则": "履约担保比例",
                        "通过": True,
                        "原文": guarantee_text,
                        "详情": f"当前比例 {percent}%，符合 ≥10% 的要求"
                    })
                else:
                    issues.append({
                        "规则": "履约担保比例",
                        "通过": False,
                        "原文": guarantee_text,
                        "详情": f"当前比例 {percent}%，小于 10%，不符合要求"
                    })
            else:
                issues.append({
                    "规则": "履约担保比例",
                    "通过": False,
                    "原文": guarantee_text,
                    "详情": "未在履约担保条款中找到明确的百分比数值"
                })
        else:
            issues.append({
                "规则": "履约担保比例",
                "通过": False,
                "原文": "未找到",
                "详情": "未找到履约担保条款"
            })
    # 规则：安全文明施工要求
    safety_text = front_info.get("安全文明施工要求", "")
    if safety_text and safety_text != "未找到":
        zero_required = [
            (r'死亡率为零|死亡率.*?0|零死亡', "未明确死亡率为零"),
            (r'重伤率为零|重伤率.*?0|零重伤', "未明确重伤率为零"),
            (r'职业病为零|职业病.*?0|零职业病', "未明确职业病为零"),
        ]
        rules_required = [
            (r'安全生产.*?管理|安全.*?规定', "未提及符合安全生产管理规定"),
            (r'文明施工.*?管理|文明.*?规定', "未提及符合文明施工管理规定"),
        ]
        missing = []
        for pattern, msg in zero_required:
            if not re.search(pattern, safety_text, re.IGNORECASE):
                missing.append(msg)
        for pattern, msg in rules_required:
            if not re.search(pattern, safety_text, re.IGNORECASE):
                missing.append(msg)
        if missing:
            issues.append({
                "规则": "安全文明施工要求",
                "通过": False,
                "原文": safety_text,
                "详情": "内容不合理：" + "；".join(missing)
            })
        else:
            issues.append({
                "规则": "安全文明施工要求",
                "通过": True,
                "原文": safety_text,
                "详情": "内容合理，包含零事故指标及安全文明管理规定"
            })
    else:
        issues.append({
            "规则": "安全文明施工要求",
            "通过": False,
            "原文": "未找到",
            "详情": "未找到安全文明施工要求条款"
        })


    return issues

# ============================================
# 负面清单检查
# ============================================
def check_negative_list(front_info, table_dict):
    issues = []
    if front_info is None:
        front_info = {}
    if table_dict is None:
        table_dict = {}
    skip_table = (table_dict == {})
    full_text = st.session_state.get("full_text", "")

    
    # ---------- 2. 资格要求禁止性条款（四个独立规则） ----------
    if not full_text:
        issues.append({
            "规则": "资格要求禁止性条款（全文扫描）",
            "通过": False,
            "原文": "未获取到全文",
            "详情": "无法进行全文扫描，请重新上传文件"
        })
    else:
        rules = [
            {
                "name": "限定营业执照经营范围",
                "groups": [
                    r'营业执照|经营范围',
                    r'必须|须|应当|要求包含|限定',
                    r'制作|销售|产品|服务'
                ],
                "require_first": True,
                "min_matches": 2,
                "desc": "限定营业执照经营范围（如要求涵盖特定产品）",
                "keywords": ['营业执照', '经营范围', '必须', '须', '应当', '要求包含', '限定', '制作', '销售', '产品', '服务']
            },
            {
                "name": "限定或指定特定专利、商标、品牌、原产地或供应商",
                "groups": [
                    r'限定|指定|使用|持有|拥有',
                    r'专利|商标|品牌|原产地|供应商'
                ],
                "min_matches": 2,
                "desc": "限定或指定特定专利、商标、品牌、原产地或供应商",
                "keywords": ['限定', '指定', '使用', '持有', '拥有', '专利', '商标', '品牌', '原产地', '供应商']
            },
            {
                "name": "要求投标人注册资本金",
                "groups": [
                    r'(?=.*(?:注册资本金|注册资本|注册资金))(?=.*(?:\d+(?:\.\d+)?\s*(?:万|元|万元以上)|必须|须|至少|不低于)).*'
                ],
                "min_matches": 1,
                "desc": "要求投标人注册资本金（同时包含注册资本词与金额或强制性表述）",
                "keywords": ['注册资本金', '注册资本', '注册资金', '万', '元', '必须', '须', '至少', '不低于']
            },
            {
                "name": "度身定向招标或不合理条件限制、排斥潜在投标人",
                "groups": [
                    r'度身定向|量身定做|不合理条件|排斥.*?潜在投标人|限制.*?公平竞争'
                ],
                "min_matches": 1,
                "desc": "度身定向招标或不合理条件限制、排斥潜在投标人",
                "keywords": ['度身定向', '量身定做', '不合理条件', '排斥', '潜在投标人', '限制', '公平竞争']
            }
        ]

        def check_rule(text, rule):
            sentences = re.split(r'[。！；\n]+', text)
            violations = []
            groups = rule["groups"]
            require_first = rule.get("require_first", False)
            min_matches = rule.get("min_matches", len(groups))
            desc = rule["desc"]
            keywords = rule.get("keywords", [])
            for sentence in sentences:
                if require_first:
                    first_match = re.search(groups[0], sentence, re.IGNORECASE)
                    if not first_match:
                        continue
                    first_pos = first_match.start()
                    matched_count = 1
                    for pattern in groups[1:]:
                        m = re.search(pattern, sentence, re.IGNORECASE)
                        if m and m.start() > first_pos:
                            matched_count += 1
                    if matched_count >= min_matches:
                        highlighted = sentence[:200]
                        for kw in keywords:
                            pat = re.compile(r'(' + re.escape(kw) + r')', re.IGNORECASE)
                            highlighted = pat.sub(r'<span style="color:red; font-weight:bold;">\1</span>', highlighted)
                        violations.append((desc, highlighted))
                else:
                    matched_count = 0
                    for pattern in groups:
                        if re.search(pattern, sentence, re.IGNORECASE):
                            matched_count += 1
                    if matched_count >= min_matches:
                        highlighted = sentence[:200]
                        for kw in keywords:
                            pat = re.compile(r'(' + re.escape(kw) + r')', re.IGNORECASE)
                            highlighted = pat.sub(r'<span style="color:red; font-weight:bold;">\1</span>', highlighted)
                        violations.append((desc, highlighted))
            return violations

        for rule in rules:
            violations = check_rule(full_text, rule)
            if violations:
                details = []
                full_original = []
                for idx, (desc, snippet) in enumerate(violations, 1):
                    details.append(f"{idx}. {desc}")
                    full_original.append(f"【违规{idx} - {desc}】<br>原文：{snippet}")
                issues.append({
                    "规则": rule["name"],
                    "通过": False,
                    "原文": "<br><br>".join(full_original),
                    "详情": "<br>".join(details)
                })
            else:
                issues.append({
                    "规则": rule["name"],
                    "通过": True,
                    "原文": "全文未发现明显的违规",
                    "详情": "未发现该类型禁止性条款"
                })


    # ---------- 新增规则：劳务分包资质检查 ----------
    contract_type = table_dict.get("承包方式", "")
    
    if "劳务分包" in contract_type:
        qual_text = front_info.get("分供商资格要求", "")
        if qual_text and qual_text != "未找到":
            # 先用字典快速匹配
            has_qualification = False
            matched_category = None
            for cat in QUALIFICATION_STANDARDS.keys():
                if cat in qual_text:
                    has_qualification = True
                    matched_category = cat
                    break
            
            # 检查是否有资质等级
            has_level = re.search(r'(?:一级|二级|三级|壹级|贰级|叁级)', qual_text)
            
            # 如果字典没匹配到，但可能有笔误或漏写的资质，用AI辅助判断
            if not has_qualification:
                ai_has = ai_extract_in_area(
                    qual_text,
                    "资质判断",
                    "判断以下文本中是否包含建筑业企业资质类别（如'建筑装修装饰工程专业承包'、'建筑工程施工总承包'、'钢结构工程专业承包'等）。"
                    "注意：即使资质名称与标准名称略有差异（如'建筑装饰装修工程承包'），也视为包含资质类别。"
                    "只输出'有'或'无'。",
                    "资格要求"
                )
                if "有" in ai_has:
                    has_qualification = True
                    matched_category = "AI识别到资质类别"
            
            if has_qualification and has_level:
                issues.append({
                    "规则": "劳务分包招标是否要求资质具备级别",
                    "通过": False,
                    "原文": qual_text[:500],
                    "详情": f"劳务分包不应要求资质类别和等级，但资格要求中包含了资质要求（{matched_category}）和资质等级"
                })
            else:
                issues.append({
                    "规则": "劳务分包招标是否要求资质具备级别",
                    "通过": True,
                    "原文": qual_text[:500],
                    "详情": "劳务分包未要求资质类别和等级，符合规定"
                })
        else:
            issues.append({
                "规则": "劳务分包招标是否要求资质具备级别",
                "通过": True,
                "原文": "未找到资格要求",
                "详情": "未找到资格要求，无法判断"
            })
    else:
        issues.append({
            "规则": "劳务分包招标是否要求资质具备级别",
            "通过": True,
            "原文": f"承包方式为：{contract_type}",
            "详情": "非劳务分包，不适用此规则"
        })
    # ---------- 调试：提取资质等级、建设规模、最高报价限价 ----------


# 原有的提取资质函数（保持不变）
    def extract_qualification(text):
        match = re.search(r'(?:具备|有)\s*([\u4e00-\u9fff]+?(?:施工总承包|专业承包|劳务分包|工程承包))\s*((?:特级|一级|二级|三级)(?:及以上|及以下)?)', text)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        match = re.search(r'([\u4e00-\u9fff]+?(?:施工总承包|专业承包|劳务分包|工程承包))\s*((?:特级|一级|二级|三级)(?:及以上|及以下)?)', text)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return "", ""

# 新增：从文本中提取价格（支持多种关键词）
    def extract_price(text):
        if not text:
            return ""
        text_clean = re.sub(r'[，,]', '', text)
        # 关键词包含“采购控制总价”等
        keywords = [
            "采购控制总价", "采购控制价", "控制总价", "最高限价",
            "控制价", "预算金额", "最高投标限价", "招标控制价"
        ]
        # 匹配关键词后的数字和单位（捕获“元”或“万元”）
        pattern = r'(?:' + '|'.join(keywords) + r')\s*[：:]\s*([\d,，.]+)\s*(万?元)'
        match = re.search(pattern, text)
        if match:
            price = match.group(1).strip()
            unit = match.group(2).strip()
        # 去除千分位逗号
            price = re.sub(r'[，,]', '', price)
            return f"{price}{unit}"
    
    # 🔧 修改：通用匹配，支持千分位逗号
        match = re.search(r'([\d,，]+\.?\d*)\s*(万?元)', text)
        if match:
            price = match.group(1).strip()
            unit = match.group(2).strip()
            price = re.sub(r'[，,]', '', price)
            return f"{price}{unit}"
        # 最后尝试只匹配数字（无单位），默认加“元”
        match = re.search(r'(\d+[,，]?\d*\.?\d*)', text)
        if match:
            price = match.group(1).strip()
            price = re.sub(r'[，,]', '', price)
            return f"{price}元"
        return ""

    # ===== 原代码中的获取部分 =====
    front_qual = front_info.get("分供商资格要求", "")
    table_qual = table_dict.get("响应人资格要求", "")

    #===== 建设规模：提取“采购项目概况”完整内容并截断 =====
    full_text = st.session_state.get("full_text", "")
    project_scale = "未找到"
    if full_text:
        # 优先级1：提取“采购项目概况”
        pattern = r'采购项目概况[：:]\s*(.*?)(?=\n\s*\n|\n\s*\d+[、\.]|\n[^\n]*[：:]|$)'
        match = re.search(pattern, full_text, re.DOTALL)
        if match:
            project_scale = match.group(1).strip()
            project_scale = re.sub(r'\s+', ' ', project_scale)
        else:
            # 优先级2：如果没有“采购项目概况”，尝试“建设规模”或“项目概况”
            fallback_match = re.search(
                r'(?:建设规模|项目概况|项目规模|工程规模)[：:]\s*(.*?)(?=\n\s*\n|\n\s*\d+[、\.]|\n[^\n]*[：:]|$)',
                full_text, re.DOTALL
            )
            if fallback_match:
                project_scale = fallback_match.group(1).strip()
                project_scale = re.sub(r'\s+', ' ', project_scale)

        # 如果找到内容，进行截断（去除“采购范围”等后续部分）
        if project_scale != "未找到":
            truncate_markers = ['2 采购范围', '二、采购范围', '2、采购范围', '二、采购范围']
            for marker in truncate_markers:
                idx = project_scale.find(marker)
                if idx != -1:
                    project_scale = project_scale[:idx].strip()
                    break
    else:
        # 若 full_text 为空，则回退到 front_info（仅作兜底）
        project_scale = front_info.get("建设规模", "未找到")


    # ---- 采购控制价提取（仅在须知表中查找） ----
    # 定义可能的关键词列表（按优先级排序）
    price_keys = [
        "采购控制价", 
        "最高报价限价", 
        "报价最高限价", 
        "最高限价", 
        "控制价", 
        "最高投标限价"
    ]
    max_price = ""
    for key in price_keys:
        if key in table_dict:
            val = table_dict.get(key, "")
            if val and val != "未找到" and val.strip():
                # 使用正则提取数字（含小数，支持千分位逗号）
                val_clean = re.sub(r'[，,\s]', '', val)  # 1,870,000.00元 → 1870000.00元
                match = re.search(r'(\d+\.?\d*)', val_clean)
                if match:
                    raw_num = match.group(1)
                    # 清理格式：去除中文逗号、英文逗号、空格
                    clean_num = re.sub(r'[，,\s]', '', raw_num)
                    max_price = clean_num
                    break
    # 如果没找到，保留空字符串（后续显示为“未提供”）
    # =================================

    # 资质提取（不变）
    front_level = extract_qualification(front_qual)
    table_level = extract_qualification(table_qual)

 

    # ---- 新增：资质等级与承包范围匹配检查（使用 Ollama） ----
    # 使用原有 extract_qualification 提取 (类别, 等级文本)
    raw_category, raw_level = extract_qualification(front_qual)
    if not raw_category or not raw_level:
        raw_category, raw_level = extract_qualification(table_qual)

    category = None
    level = None
    match_note = ""  # 用于记录匹配说明

    if raw_category and raw_level:
        # 1. 映射类别：使用 rapidfuzz 在字典中寻找最相似的键
        # 1. 映射类别：使用 rapidfuzz 在字典中寻找最相似的键
        from rapidfuzz import fuzz
        best_score = 0
        best_match = None
        for std_cat in QUALIFICATION_STANDARDS.keys():
            score = fuzz.ratio(raw_category, std_cat)
            if score > best_score:
                best_score = score
                best_match = std_cat
        # 如果最佳匹配度 >= 70，则使用该匹配；否则保留原始名称
        if best_score >= 70:
            category = best_match
            # 仅当相似度 < 100% 时才生成注释，否则注释为空
            if best_score < 100:
                match_note = f"（⚠️ 原始提取为“{raw_category}”，匹配到字典键“{best_match}”，相似度{best_score}%）"
            else:
                match_note = ""  # 完全匹配，不显示注释
        else:
            category = raw_category
            match_note = f"（⚠️ 未在字典中找到相似键）"


        # 2. 标准化等级：提取“一级/二级/三级”基础等级
        level_match = re.search(r'(一级|二级|三级|壹级|贰级|叁级)', raw_level)
        if level_match:
            level_map = {
                "一级": "一级", "壹级": "一级",
                "二级": "二级", "贰级": "二级",
                "三级": "三级", "叁级": "三级"
            }
            level = level_map.get(level_match.group(1))
        # 若没匹配到，尝试从“二级及以上”中提取“二级”
        if not level:
            level_match = re.search(r'(一级|二级|三级)\s*(?:及|含|以上|以下)', raw_level)
            if level_match:
                level = level_match.group(1)

    # 若成功提取到 category 和 level，继续检查
    if category and level:
        desc = None
        if category in QUALIFICATION_STANDARDS:
            desc = QUALIFICATION_STANDARDS[category].get(level)
        if desc:
            price_str = max_price if max_price else "未提供"
            scale_str = project_scale if project_scale else "未提供"
            is_ok, reason = check_with_ollama(category, level, desc, price_str, scale_str)

            # ---- 构建显示信息 ----
            scale_display = project_scale if project_scale != "未找到" else "未提供"
            if match_note:
                qual_display = f"资质：{category} {level} {match_note}"
            else:
                qual_display = f"资质：{category} {level}"
            original_text = (
                f"{qual_display}\n"
                f"标准：{desc}\n"
                f"建设规模：{scale_display}\n"
                f"采购控制价: {max_price}\n"
            )

            issues.append({
                "规则": "资质等级标准合理性",
                "通过": is_ok,
                "原文": original_text,
                "详情": reason
            })
        else:
            # category 和 level 都有，但 desc 为空（等级不存在）
            if category in QUALIFICATION_STANDARDS:
                available_levels = list(QUALIFICATION_STANDARDS[category].keys())
                available_levels_str = "、".join(available_levels)
                issues.append({
                    "规则": "资质等级标准合理性",
                    "通过": False,
                    "原文": f"提取到资质 {category} {level}，该类别存在但字典中无对应等级",
                    "详情": f"【等级不存在】{category} 支持的等级有：{available_levels_str}，但提取到的等级是 {level}，请检查是否提取正确"
                })
            else:
                issues.append({
                    "规则": "资质等级标准合理性",
                    "通过": False,
                    "原文": f"提取到资质 {category} {level}，但字典中无该类别",
                    "详情": "【类别不存在】请检查《建筑业企业资质标准》是否包含该类别"
                })
    else:
        # category 或 level 为空
        if category in QUALIFICATION_STANDARDS:
            # 有 category，但 level 为空
            available_levels = list(QUALIFICATION_STANDARDS[category].keys())
            available_levels_str = "、".join(available_levels)
            issues.append({
                "规则": "资质等级标准合理性",
                "通过": False,
                "原文": f"提取到资质类别 {category}，但未提取到等级",
                "详情": f"【等级缺失】{category} 支持的等级有：{available_levels_str}，但未提取到具体等级，请检查资格要求文本"
            })
        else:
            issues.append({
                "规则": "资质等级标准合理性",
                "通过": False,
                "原文": f"未提取到有效的资质类别或等级（原始提取结果：类别={raw_category}, 等级={raw_level}）",
                "详情": "请检查资格要求文本"
            })

    # 3. 质量保证金比例检查（全文扫描）
    if full_text:
        pattern = r'(?:工程质量保证金|质量保证金)\s*[为:：]?\s*(\d+(?:\.\d+)?)\s*%'
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            percent = float(match.group(1))
            start = max(0, match.start() - 50)
            end = min(len(full_text), match.end() + 200)
            context = full_text[start:end].replace('\n', ' ')
            if percent > 3:
                issues.append({
                    "规则": "质量保证金比例",
                    "通过": False,
                    "原文": context[:300],
                    "详情": f"质量保证金比例为 {percent}%，大于 3%，不符合要求"
                })
            else:
                issues.append({
                    "规则": "质量保证金比例",
                    "通过": True,
                    "原文": context[:300],
                    "详情": f"质量保证金比例为 {percent}%，符合 ≤3% 的要求"
                })
        else:
            sentences = re.split(r'[。！；\n]+', full_text)
            found = False
            for sentence in sentences:
                if ('质保金' in sentence or '质量保证金' in sentence) and '%' in sentence:
                    found = True
                    percent_match = re.search(r'(\d+(?:\.\d+)?)\s*%', sentence)
                    if percent_match:
                        percent = float(percent_match.group(1))
                        if percent > 3:
                            issues.append({
                                "规则": "质量保证金比例",
                                "通过": False,
                                "原文": sentence[:200],
                                "详情": f"质量保证金比例为 {percent}%，大于 3%，不符合要求"
                            })
                        else:
                            issues.append({
                                "规则": "质量保证金比例",
                                "通过": True,
                                "原文": sentence[:200],
                                "详情": f"质量保证金比例为 {percent}%，符合 ≤3% 的要求"
                            })
                    else:
                        issues.append({
                            "规则": "质量保证金比例",
                            "通过": False,
                            "原文": sentence[:200],
                            "详情": "未找到明确的百分比数值"
                        })
                    break
            if not found:
                issues.append({
                    "规则": "质量保证金比例",
                    "通过": False,
                    "原文": "未找到",
                    "详情": "未找到质量保证金或质保金条款"
                })
    else:
        issues.append({
            "规则": "质量保证金比例",
            "通过": False,
            "原文": "未获取到全文",
            "详情": "无法进行全文扫描，请重新上传文件"
        })

    # 4. 投标保证金比例/金额限制检查（全文扫描）
    if full_text:
        sentences = re.split(r'[。！；\n]+', full_text)
        found = False
        for sentence in sentences:
            if '投标保证金' in sentence:
                found = True
                percent_match = re.search(r'(\d+(?:\.\d+)?)\s*%', sentence)
                if percent_match:
                    percent = float(percent_match.group(1))
                    if percent > 2:
                        issues.append({
                            "规则": "投标保证金比例",
                            "通过": False,
                            "原文": sentence[:200],
                            "详情": f"投标保证金比例 {percent}%，超过法定上限 2%"
                        })
                    else:
                        issues.append({
                            "规则": "投标保证金比例",
                            "通过": True,
                            "原文": sentence[:200],
                            "详情": f"投标保证金比例 {percent}%，符合 ≤2% 的要求"
                        })
                amount_match = re.search(r'(\d+(?:\.\d+)?)\s*(万?元?|万元)', sentence)
                if amount_match:
                    num = float(amount_match.group(1))
                    unit = amount_match.group(2)
                    if '万' in unit:
                        amount = num * 10000
                    else:
                        amount = num
                    if amount > 800000:
                        issues.append({
                            "规则": "投标保证金金额",
                            "通过": False,
                            "原文": sentence[:200],
                            "详情": f"投标保证金金额 {amount/10000:.0f}万元，超过法定上限 80万元"
                        })
                    else:
                        issues.append({
                            "规则": "投标保证金金额",
                            "通过": True,
                            "原文": sentence[:200],
                            "详情": f"投标保证金金额 {amount/10000:.0f}万元，符合 ≤80万元 的要求"
                        })
                if percent_match is None and amount_match is None:
                    issues.append({
                        "规则": "投标保证金条款",
                        "通过": False,
                        "原文": sentence[:200],
                        "详情": "未找到明确的百分比或金额数字"
                    })
                break
        if not found:
            issues.append({
                "规则": "投标保证金条款",
                "通过": False,
                "原文": "未找到",
                "详情": "未找到投标保证金相关条款"
            })
    else:
        issues.append({
            "规则": "投标保证金条款",
            "通过": False,
            "原文": "未获取到全文",
            "详情": "无法进行全文扫描，请重新上传文件"
        })

    # ---------- 5. 评标方式（仅在未跳过表格时检查） ----------

    if not skip_table:
        judge_key = None
        for k in table_dict.keys():
            # 扩展匹配关键词：评委、评标、评审、委员会、谈判小组、成员人数等
            if any(word in k for word in ['评委', '评标', '评审', '委员会', '谈判', '成员']) and any(word in k for word in ['人数', '成员', '组成', '小组']):
                judge_key = k
                break
        if judge_key:
            judge_text = table_dict[judge_key]
            
            # 第一步：尝试提取具体人数
            numbers = re.findall(r'(\d+)', judge_text)
            valid_numbers = [int(n) for n in numbers if 5 <= int(n) <= 50]
            
            if valid_numbers:
                # 有具体人数，直接判断
                num = valid_numbers[0]  # 取第一个
                if num >= 5 and num % 2 == 1:
                    issues.append({
                        "规则": "评标方式",
                        "通过": True,
                        "原文": judge_text,
                        "详情": f"评委人数为 {num} 人，符合5人及以上单数的要求"
                    })
                else:
                    issues.append({
                        "规则": "评标方式",
                        "通过": False,
                        "原文": judge_text,
                        "详情": f"评委人数为 {num} 人，{'小于5人' if num < 5 else '为偶数，不符合单数要求'}"
                    })
            else:
                # 第二步：没有具体人数，检查是否有"单数"或"奇数"字样
                has_odd = re.search(r'单数|奇数', judge_text)
                if has_odd:
                    # 有字样但没有数字，无法确认具体人数，视为不通过（因为可能是3人）
                    issues.append({
                        "规则": "评标方式",
                        "通过": False,
                        "原文": judge_text,
                        "详情": "未明确写明具体人数，仅提及单数/奇数要求，无法确认是否达到5人及以上"
                    })
                else:
                    issues.append({
                        "规则": "评标方式",
                        "通过": False,
                        "原文": judge_text,
                        "详情": "未明确写明评委人数或单数要求"
                    })
        else:
            issues.append({
                "规则": "评标方式",
                "通过": False,
                "原文": "未找到",
                "详情": "须知表中未找到评标方式相关条款"
            })

    #6 ---------- 新增规则：通用投标时间过短检查（所有合同，小于5天即违规） ----------
    full_text = st.session_state.get("full_text", "")
    if full_text:
        issue_date = front_info.get("发文时间", "")
        deadline_str = front_info.get("截止时间", "")
        if issue_date != "未找到" and deadline_str != "未找到":
            try:
                issue_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', issue_date)
                if issue_match:
                    iy, im, iday = issue_match.groups()
                    issue_clean = f"{iy}年{int(im):02d}月{int(iday):02d}日"
                    deadline_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', deadline_str)
                    if deadline_match:
                        dy, dm, dday = deadline_match.groups()
                        deadline_clean = f"{dy}年{int(dm):02d}月{int(dday):02d}日"
                        start = datetime.strptime(issue_clean, "%Y年%m月%d日")
                        end = datetime.strptime(deadline_clean, "%Y年%m月%d日")
                        days = (end - start).days + 1  # 包含起止日
                        if days < 5:
                            issues.append({
                                "规则": "招标文件发售期不得少于5日",
                                "通过": False,
                                "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}",
                                "详情": f"招标发售期仅 {days} 日，低于最小限制5日"
                            })
                        else:
                            issues.append({
                                "规则": "招标文件发售期不得少于5日",
                                "通过": True,
                                "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}",
                                "详情": f"招标发售期 {days} 日，满足最小限制5日"
                            })
            except Exception:
                pass  # 日期解析失败则跳过此规则



    #7 ---------- 规则：依法必招项目投标截止日期检查 ----------

    full_text = st.session_state.get("full_text", "")
    if full_text:
        procurement_method = table_dict.get("采购方式", "")
        if procurement_method == "" or procurement_method == "未找到":
            procurement_method = front_info.get("采购方式", "")
        issue_date = front_info.get("发文时间", "")
        deadline_str = front_info.get("截止时间", "")
        
        if issue_date != "未找到" and deadline_str != "未找到" and procurement_method != "未找到" and procurement_method != "":
            try:
                issue_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', issue_date)
                if not issue_match:
                    raise ValueError("发文时间格式错误")
                iy, im, iday = issue_match.groups()
                issue_clean = f"{iy}年{int(im):02d}月{int(iday):02d}日"
                
                deadline_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', deadline_str)
                if not deadline_match:
                    raise ValueError("截止时间未找到日期部分")
                dy, dm, dday = deadline_match.groups()
                deadline_clean = f"{dy}年{int(dm):02d}月{int(dday):02d}日"
                
                start = datetime.strptime(issue_clean, "%Y年%m月%d日")
                end = datetime.strptime(deadline_clean, "%Y年%m月%d日")
                days = (end - start).days + 1
                
                # 判断是否为依法必招
                if "公开招标" in procurement_method and "依法必招" in procurement_method:
                    required_days = 20
                    if days < required_days:
                        issues.append({
                            "规则": "依法必招投标截止日期",
                            "通过": False,
                            "原文": f"采购方式：{procurement_method}，发文时间：{issue_date}，截止时间：{deadline_str}",
                            "详情": f"依法必招项目招标文件开始发出之日起至投标人提交投标文件截止之日仅 {days} 天，低于法定最低要求 {required_days} 天"
                        })
                    else:
                        issues.append({
                            "规则": "依法必招投标截止日期",
                            "通过": True,
                            "原文": f"采购方式：{procurement_method}，发文时间：{issue_date}，截止时间：{deadline_str}",
                            "详情": f"依法必招项目招标文件开始发出之日起至投标人提交投标文件截止之日仅 {days} 天，满足最低要求 {required_days} 天"
                        })
                else:
                    # 不是依法必招或无法判断，直接通过
                    issues.append({
                        "规则": "依法必招投标截止日期",
                        "通过": True,
                        "原文": f"采购方式：{procurement_method}",
                        "详情": "非依法必招项目或未明确采购方式，不进行天数检查"
                    })
            except Exception as e:
                issues.append({
                    "规则": "依法必招投标截止日期",
                    "通过": False,
                    "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}",
                    "详情": f"日期解析失败：{str(e)}"
                })
        else:
            missing = []
            if issue_date == "未找到":
                missing.append("发文时间")
            if deadline_str == "未找到":
                missing.append("截止时间")
            if not procurement_method or procurement_method == "未找到":
                missing.append("采购方式")
            if missing:
                issues.append({
                    "规则": "依法必招投标截止日期",
                    "通过": False,
                    "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}，采购方式：{procurement_method}",
                    "详情": f"缺少必要字段：{', '.join(missing)}"
                })

    #8 ---------- 规则：非依法必招项目投标截止日期检查（基于分包类型） ----------
    if full_text:
        procurement_method = table_dict.get("采购方式", "")
        if procurement_method == "" or procurement_method == "未找到":
            procurement_method = front_info.get("采购方式", "")
        issue_date = front_info.get("发文时间", "")
        deadline_str = front_info.get("截止时间", "")
        
        if issue_date != "未找到" and deadline_str != "未找到" and procurement_method != "未找到" and procurement_method != "":
            # 仅对非依法必招项目（或不包含“依法必招”的公开招标）进行检查
            if not ("公开招标" in procurement_method and "依法必招" in procurement_method):
                try:
                    issue_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', issue_date)
                    if not issue_match:
                        raise ValueError("发文时间格式错误")
                    iy, im, iday = issue_match.groups()
                    issue_clean = f"{iy}年{int(im):02d}月{int(iday):02d}日"
                    
                    deadline_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', deadline_str)
                    if not deadline_match:
                        raise ValueError("截止时间未找到日期部分")
                    dy, dm, dday = deadline_match.groups()
                    deadline_clean = f"{dy}年{int(dm):02d}月{int(dday):02d}日"
                    
                    start = datetime.strptime(issue_clean, "%Y年%m月%d日")
                    end = datetime.strptime(deadline_clean, "%Y年%m月%d日")
                    days = (end - start).days + 1
                    
                    # 确定分包类型
                    contract_type = None
                    contract_method = table_dict.get("承包方式", "")
                    if contract_method and contract_method != "未找到":
                        if "专业分包" in contract_method:
                            contract_type = "专业分包"
                        elif "劳务分包" in contract_method:
                            contract_type = "劳务分包"
                        elif "材料采购" in contract_method:
                            contract_type = "材料采购"
                        elif "服务分包" in contract_method:
                            contract_type = "服务分包"
                        elif "机械租赁" in contract_method:
                            contract_type = "机械租赁"
                        elif "设备采购" in contract_method:
                            contract_type = "设备采购"
                    
                    if contract_type is None:
                        project_name = front_info.get("采购项目名称", "")
                        if "专业分包" in project_name:
                            contract_type = "专业分包"
                        elif "劳务分包" in project_name:
                            contract_type = "劳务分包"
                        elif "材料采购" in project_name:
                            contract_type = "材料采购"
                        elif "服务分包" in project_name:
                            contract_type = "服务分包"
                        elif "机械租赁" in project_name:
                            contract_type = "机械租赁"
                        elif "设备采购" in project_name:
                            contract_type = "设备采购"
                        else:
                            if "专业" in project_name:
                                contract_type = "专业分包"
                            elif "劳务" in project_name:
                                contract_type = "劳务分包"
                            elif "材料" in project_name:
                                contract_type = "材料采购"
                            elif "服务" in project_name:
                                contract_type = "服务分包"
                            elif "租赁" in project_name:
                                contract_type = "机械租赁"
                            elif "设备" in project_name:
                                contract_type = "设备采购"
                    
                    if contract_type in ["专业分包", "设备采购"]:
                        required_days = 10
                        check_type = "专业分包/设备采购"
                    elif contract_type in ["劳务分包", "材料采购"]:
                        required_days = 7
                        check_type = "劳务分包/材料采购"
                    else:
                        issues.append({
                            "规则": "非依法必招投标截止日期",
                            "通过": False,
                            "原文": f"承包方式：{contract_method}，采购项目名称：{front_info.get('采购项目名称', '')}",
                            "详情": "未找到有效的承包方式（专业分包/劳务分包/材料采购/服务分包/机械租赁/设备采购），无法判断投标截止日期要求"
                        })
                        required_days = None
                    
                    if required_days is not None:
                        if days < required_days:
                            issues.append({
                                "规则": "非依法必招投标截止日期",
                                "通过": False,
                                "原文": f"分包类型：{check_type}，发文时间：{issue_date}，截止时间：{deadline_str}",
                                "详情": f"招标文件开始发出之日起至投标人提交投标文件截止之日仅 {days} 天，低于法定最低要求 {required_days} 天"
                            })
                        else:
                            issues.append({
                                "规则": "非依法必招投标截止日期",
                                "通过": True,
                                "原文": f"分包类型：{check_type}，发文时间：{issue_date}，截止时间：{deadline_str}",
                                "详情": f"招标文件开始发出之日起至投标人提交投标文件截止之日 {days} 天，满足最低要求 {required_days} 天"
                            })
                except Exception as e:
                    issues.append({
                        "规则": "非依法必招投标截止日期",
                        "通过": False,
                        "原文": f"发文时间：{issue_date}，截止时间：{deadline_str}",
                        "详情": f"日期解析失败：{str(e)}"
                    })
        else:
            # 如果缺少必要字段，为避免重复提示，可以添加一条综合提示（可选）
            pass

#9 开标日期与投标截止日期一致
    full_text = st.session_state.get("full_text", "")

    if full_text:
        same_time_match = re.search(
            r'响应文件开启在响应文件递交截止时间的同一时间进行',
            full_text,
            re.DOTALL | re.IGNORECASE
        )

        if same_time_match:
            issues.append({
                "规则": "投标截止时间与开标时间一致性",
                "通过": True,
                "原文": same_time_match.group(0),
                "详情": "文件已明确响应文件开启在递交截止时间的同一时间进行，符合法规要求"
            })
        else:
            deadline = table_dict.get("截止时间", "") or front_info.get("截止时间", "")
            opening_time = ""
            patterns = [
                r'响应文件开启时间[（(]?(?:开标时间)?[）)]?[为]?\s*(.+?)(?:，|；|。|\n|$)',
                r'开标时间[：:]\s*(.+?)(?:，|；|。|\n|$)',
                r'开启时间[：:]\s*(.+?)(?:，|；|。|\n|$)',
            ]
            for pat in patterns:
                match = re.search(pat, full_text, re.DOTALL)
                if match:
                    opening_time = match.group(1).strip()
                    break

            def extract_date(text):
                match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', text)
                if match:
                    return f"{match.group(1)}年{int(match.group(2)):02d}月{int(match.group(3)):02d}日"
                return None

            if deadline and deadline != "未找到" and opening_time and opening_time != "未找到":
                deadline_date = extract_date(deadline)
                opening_date = extract_date(opening_time)

                if deadline_date and opening_date:
                    if deadline_date == opening_date:
                        issues.append({
                            "规则": "投标截止时间与开标时间一致性",
                            "通过": True,
                            "原文": f"截止时间：{deadline}\n开标时间：{opening_time}",
                            "详情": f"投标截止时间与开标时间均为 {deadline_date}，一致"
                        })
                    else:
                        issues.append({
                            "规则": "投标截止时间与开标时间一致性",
                            "通过": False,
                            "原文": f"截止时间：{deadline}\n开标时间：{opening_time}",
                            "详情": f"投标截止时间（{deadline_date}）与开标时间（{opening_date}）不一致"
                        })
                else:
                    issues.append({
                        "规则": "投标截止时间与开标时间一致性",
                        "通过": False,
                        "原文": f"截止时间：{deadline}\n开标时间：{opening_time}",
                        "详情": "无法提取完整的日期信息进行比较"
                    })
            elif deadline and deadline != "未找到":
                issues.append({
                    "规则": "投标截止时间与开标时间一致性",
                    "通过": False,
                    "原文": f"截止时间：{deadline}",
                    "详情": "找到了投标截止时间，但未找到开标时间"
                })
            elif opening_time and opening_time != "未找到":
                issues.append({
                    "规则": "投标截止时间与开标时间一致性",
                    "通过": False,
                    "原文": f"开标时间：{opening_time}",
                    "详情": "找到了开标时间，但未找到投标截止时间"
                })
            else:
                issues.append({
                    "规则": "投标截止时间与开标时间一致性",
                    "通过": False,
                    "原文": "未找到",
                    "详情": "未找到投标截止时间和开标时间，无法检查一致性"
                })




# =================ai审查法规======================



    full_text = st.session_state.get("full_text", "")
    
    if full_text:
        # 分批规则（排除4项正则规则）
        rule_batches = [
            # 第一批：保证金类（3条）
            [
                "设立其他非法保证金",
                "无息退还投标保证金或退还时间不合规",
                "以不合理的理由没收投标保证金"
            ],
            # 第二批：劳务分包类（3条）
            [
                "劳务分包招标文件中包含主材、建筑材料款、机械费、周转材料等",
                "劳务分包招标文件原封不动引用主合同清单",
                "分包招标文件编制中包含项目主体结构部分的清单"
            ],
            # 第三批：商务价格+招标程序（5条）
            [
                "招标文件中规定单价与数量乘积与合价不一致时按就低不就高原则修正",
                "招标文件中设置开标谈判程序",
                "开标后出现第二轮竞价或第二次谈判条款",
                "无正当理由随意暂停或终止招标",
                "要求不合理的签订合同时限"
            ]
        ]
        
        for batch in rule_batches:
            rules_to_check = []
            for check in COMPLIANCE_CHECKS:
                if check["name"] in batch:
                    rules_to_check.append({
                        "name": check["name"],
                        "description": check["description"],
                        "legal_basis": check["legal_basis"]
                    })
            
            if rules_to_check:
                ai_results = ai_rules_check(full_text, rules_to_check)
                
                if ai_results and isinstance(ai_results, list) and len(ai_results) > 0:
                    for result in ai_results:
                        if isinstance(result, dict) and "rule" in result:
                            rule_name_from_ai = result.get("rule", "")
                            
                            matched_rule = None
                            for check in COMPLIANCE_CHECKS:
                                if check["name"] in rule_name_from_ai or rule_name_from_ai in check["name"]:
                                    matched_rule = check["name"]
                                    break
                            
                            rule_name = matched_rule if matched_rule else rule_name_from_ai
                            violation = result.get("violation", False)
                            evidence = result.get("evidence", "")
                            reason = result.get("reason", "")
                            
                            issues.append({
                                "规则": rule_name,
                                "通过": not violation,
                                "原文": evidence if evidence and evidence != "无" else "未发现违规",
                                "详情": reason if reason else "AI审查完成"
                            })
                        else:
                            print(f"⚠️ 跳过无效结果: {result}")
                else:
                    print(f"⚠️ 批次 {batch} 审查返回空结果或无有效数据")
        
        return issues

# ============================================
# Streamlit UI
# ============================================
st.set_page_config(page_title="招标文件审查助手", layout="wide")
st.title("招标文件审查助手")



file_type = st.selectbox("请选择文件类型", ["", "招标文件", "响应文件"], index=0)

if file_type == "":
    st.warning("请先选择文件类型")
    st.stop()

if file_type == "招标文件":
    uploaded_file = st.file_uploader("上传文件 (.docx / .pdf / .txt)", type=["docx", "pdf", "txt"])

    # 初始化 session_state
    if "front_info" not in st.session_state:
        st.session_state.front_info = None
    if "table_dict" not in st.session_state:
        st.session_state.table_dict = None
    if "results" not in st.session_state:
        st.session_state.results = None
    if "issues" not in st.session_state:
        st.session_state.issues = None
    if "last_file" not in st.session_state:
        st.session_state.last_file = None
    if "full_text" not in st.session_state:
        st.session_state.full_text = None
    if "negative_issues" not in st.session_state:
        st.session_state.negative_issues = None
    if "skip_table" not in st.session_state:
        st.session_state.skip_table = False
    if "front_info_ai_only" not in st.session_state:
        st.session_state.front_info_ai_only = None

    # 上传文件处理
    if uploaded_file is not None and st.session_state.last_file != uploaded_file.name:
        with st.spinner("正在提取公告信息..."):
            full_text = extract_contract_text(uploaded_file)
            st.session_state.full_text = full_text
            cutoff = full_text.find("须知样表")
            if cutoff != -1:
                front_text = full_text[:cutoff]
            else:
                front_text = full_text
            st.session_state.front_info = parse_front_info_ai_only(front_text)
            st.session_state.table_dict = None
            st.session_state.results = None
            st.session_state.issues = None
            st.session_state.negative_issues = None
            st.session_state.last_file = uploaded_file.name

    # 显示公告信息
    if st.session_state.front_info:
        with st.expander("提取的公告信息", expanded=False):
            st.json(st.session_state.front_info)

    # 须知表获取
    if uploaded_file is not None:
        st.subheader("获取“须知样表”内容")
        use_auto = st.radio("表格来源", ["手动粘贴", "自动提取（仅限原生表格）"], horizontal=True)
        
        if use_auto == "手动粘贴":
            if st.session_state.table_dict is None:
                manual_text = st.text_area(
                    "请将“须知样表”的表格内容（包括表头及所有数据行）完整粘贴到下方",
                    height=300,
                    help="从 Word 文档中复制表格，粘贴到这里。请确保包含项目、内容、说明与要求等列。"
                )
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("解析表格", key="parse_manual"):
                        if manual_text.strip():
                            table_dict = parse_table_from_text(manual_text)
                            if table_dict:
                                st.session_state.table_dict = table_dict
                                st.success(f"成功解析 {len(table_dict)} 项")
                            else:
                                st.error("解析失败，请检查粘贴内容是否包含表格数据。")
                        else:
                            st.warning("请先粘贴表格内容")
                with col2:
                    if st.button("跳过表格", key="skip_table"):
                        st.session_state.table_dict = {}
                        st.success("已跳过须知表，将不进行依赖表格的检查")
            else:
                if st.session_state.table_dict == {}:
                    st.info("已跳过须知表，未提供表格内容")
                else:
                    with st.expander("须知样表内容", expanded=False):
                        st.json(st.session_state.table_dict)
        else:  # 自动提取
            if st.session_state.table_dict is None:
                with st.spinner("尝试自动提取表格..."):
                    full_text = extract_contract_text(uploaded_file)
                    lines = full_text.split('\n')
                    start_idx = None
                    for i, line in enumerate(lines):
                        if '分供商须知样表' in line:
                            start_idx = i
                            break
                    if start_idx is not None:
                        table_lines = []
                        for i in range(start_idx + 1, len(lines)):
                            line = lines[i].strip()
                            if not line:
                                break
                            if line.startswith('|') or re.match(r'^\d+\s+', line):
                                table_lines.append(line)
                            else:
                                break
                        if table_lines:
                            table_dict = parse_table_from_text("\n".join(table_lines))
                            if table_dict:
                                st.session_state.table_dict = table_dict
                                st.success("自动提取成功")
                            else:
                                st.error("自动提取失败，请切换为手动粘贴模式")
                                st.session_state.table_dict = None
                        else:
                            st.error("未找到表格数据，请切换为手动粘贴模式")
                            st.session_state.table_dict = None
                    else:
                        st.error("未找到“须知样表”标题，请切换为手动粘贴模式")
                        st.session_state.table_dict = None
                if st.session_state.table_dict is None:
                    st.warning("自动提取失败，您也可以手动粘贴表格内容：")
                    manual_fallback = st.text_area("请粘贴表格", height=200)
                    if st.button("手动解析", key="parse_fallback"):
                        if manual_fallback.strip():
                            table_dict = parse_table_from_text(manual_fallback)
                            if table_dict:
                                st.session_state.table_dict = table_dict
                                st.success("手动解析成功")
                            else:
                                st.error("解析失败")
            else:
                if st.session_state.table_dict == {}:
                    st.info("已跳过须知表，未提供表格内容")
                else:
                    with st.expander("须知样表内容", expanded=False):
                        st.json(st.session_state.table_dict)

    # 开始检查按钮
    if st.button("开始检查", type="primary"):
        if st.session_state.front_info:
            if st.session_state.table_dict == {}:
                st.warning("已跳过须知表，上下文一致性检查将不执行，其他检查照常进行。")
                st.session_state.negative_issues = check_negative_list(st.session_state.front_info, st.session_state.table_dict)
                issues = check_internal_rules(st.session_state.front_info, st.session_state.table_dict)
                st.session_state.issues = issues
                st.session_state.results = []
            else:
                st.session_state.negative_issues = check_negative_list(st.session_state.front_info, st.session_state.table_dict)
                issues = check_internal_rules(st.session_state.front_info, st.session_state.table_dict)
                st.session_state.issues = issues
                st.session_state.results = compare_with_rules(st.session_state.front_info, st.session_state.table_dict)

    # ========== 负面清单检查显示 ==========
    if st.session_state.get("negative_issues"):
        if "show_negative" not in st.session_state:
            st.session_state.show_negative = False
        col1, col2 = st.columns([0.1, 5])
        with col1:
            arrow = "▼" if st.session_state.show_negative else "▶"
            if st.button(arrow, key="negative_arrow", help="折叠/展开"):
                st.session_state.show_negative = not st.session_state.show_negative
        with col2:
            st.markdown("<h2 style='font-size: 28px; margin: 0;'>负面清单检查</h2>", unsafe_allow_html=True)
        
        if st.session_state.show_negative:
            # 定义分类规则名称
            qualification_keywords = [
                "限定营业执照经营范围",
                "限定或指定特定专利、商标、品牌、原产地或供应商",
                "要求投标人注册资本金",
                "限定营业执照经营范围（AI审查）",
                "限定或指定特定专利、商标、品牌、原产地或供应商（AI审查）",
                "要求投标人注册资本金（AI审查）",
                "度身定向招标或不合理条件限制、排斥潜在投标人",
                "劳务分包招标是否要求资质具备级别",
                "资质等级标准合理性"
            ]
            guarantee_keywords = [
                "设立其他非法保证金",
                "无息退还投标保证金或退还时间不合规",
                "以不合理的理由没收投标保证金",
                "投标保证金比例",
                "投标保证金金额",
                "投标保证金条款",
                "质量保证金比例"
            ]

            def sort_by_order(issue_list, order_list):
                """按指定顺序排序"""
                def get_order(issue):
                    rule_name = issue["规则"]
                    for i, name in enumerate(order_list):
                        if name in rule_name or rule_name in name:
                            return i
                    return len(order_list)  # 未匹配的放最后
                return sorted(issue_list, key=get_order)

            deadline_keywords = [
                "招标文件发售期",
                "非依法必招投标截止日期",
                "依法必招投标截止日期",
                "投标截止时间与开标时间一致性"
            ]
            contract_keywords = [
                "劳务分包招标文件中包含主材、建筑材料款、机械费、周转材料等",
                "劳务分包招标文件原封不动引用主合同清单",
                "分包招标文件编制中包含项目主体结构部分的清单"
            ]
            correction_keywords = [
                "招标文件中规定单价与数量乘积与合价不一致时按就低不就高原则修正"
            ]
            workflow_keywords = [
                "招标文件中设置开标谈判程序",
                "开标后出现第二轮竞价或第二次谈判条款",
                "无正当理由随意暂停或终止招标",
                "要求不合理的签订合同时限"
            ]
            
            # 分类
            qual_issues = []
            guar_issues = []
            deadline_issues = []
            contract_issues = []
            correction_issues = []
            workflow_issues = []
            other_issues = []
            
            for issue in st.session_state.negative_issues:
                rule_name = issue["规则"]
                if rule_name in qualification_keywords:
                    qual_issues.append(issue)
                elif rule_name in guarantee_keywords:
                    guar_issues.append(issue)
                elif any(kw in rule_name for kw in deadline_keywords):
                    deadline_issues.append(issue)
                elif rule_name in contract_keywords:
                    contract_issues.append(issue)
                elif rule_name in correction_keywords:
                    correction_issues.append(issue)
                elif rule_name in workflow_keywords:
                    workflow_issues.append(issue)
                else:
                    other_issues.append(issue)
            
            # ---------- 资格限定折叠块 ----------
            if qual_issues:
                with st.expander("📋 资格限定", expanded=False):
                    for issue in qual_issues:
                        is_ok = issue["通过"]  
                        reason = issue["详情"]
                        if is_ok and "所有限制均不适用" in reason:
                            st.warning(f"⚠️ {issue['规则']}：{reason}")
                        elif is_ok and "所有匹配限制均合规" in reason:
                            st.success(f"✅ {issue['规则']}：{reason}")
                        elif is_ok:
                            st.success(f"✅ {issue['规则']}：{reason}")
                        else:
                            st.error(f"❌ {issue['规则']}：{reason}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 保证金约定折叠块 ----------
            if guar_issues:
                guar_issues = sort_by_order(guar_issues, guarantee_keywords)
                with st.expander("📋 保证金约定", expanded=False):
                    for issue in guar_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 招投标期限折叠块 ----------
            if deadline_issues:
                order = ["招标文件发售期", "非依法必招投标截止日期", "依法必招投标截止日期", "投标截止时间与开标时间一致性"]
                sorted_issues = []
                for kw in order:
                    for issue in deadline_issues:
                        if kw in issue["规则"]:
                            sorted_issues.append(issue)
                            break
                for issue in deadline_issues:
                    if issue not in sorted_issues:
                        sorted_issues.append(issue)
                with st.expander("📋 招投标期限", expanded=False):
                    for issue in sorted_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 招标清单折叠块 ----------
            if contract_issues:
                with st.expander("📋 招标清单", expanded=False):
                    for issue in contract_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 商务价格折叠块 ----------
            if correction_issues:
                with st.expander("📋 商务价格", expanded=False):
                    for issue in correction_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 招标程序折叠块 ----------
            if workflow_issues:
                with st.expander("📋 招标程序", expanded=False):
                    for issue in workflow_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)
            
            # ---------- 其他规则折叠块 ----------
            if other_issues:
                with st.expander("📋 其他", expanded=False):
                    for issue in other_issues:
                        if issue["通过"]:
                            st.success(f"✅ {issue['规则']}：{issue['详情']}")
                        else:
                            st.error(f"❌ {issue['规则']}：{issue['详情']}")
                        with st.expander("查看详情"):
                            st.markdown(f"**原文**：<br>{issue['原文'].replace(chr(10), '<br>')}", unsafe_allow_html=True)

    #内容合理性检查显示
    if st.session_state.get("issues"):
        if "show_content" not in st.session_state:
            st.session_state.show_content = False
        col1, col2 = st.columns([0.1, 5])
        with col1:
            arrow = "▼" if st.session_state.show_content else "▶"
            if st.button(arrow, key="content_arrow", help="折叠/展开"):
                st.session_state.show_content = not st.session_state.show_content
        with col2:
            st.markdown("<h2 style='font-size: 28px; margin: 0;'>内容合理性检查</h2>", unsafe_allow_html=True)
        if st.session_state.show_content:
            for issue in st.session_state.issues:
                if issue["通过"]:
                    st.success(f"✅  {issue['规则']}：{issue['详情']}")
                else:
                    st.error(f"❌{issue['规则']}：{issue['详情']}")
                with st.expander("查看详情"):
                    st.write(f"**原文**：{issue['原文']}")

# 上下文一致性检查显示
    if st.session_state.table_dict == {}:
        st.info("已跳过须知表，不进行上下文一致性检查")
    elif st.session_state.get("results"):
        if "show_consistency" not in st.session_state:
            st.session_state.show_consistency = False
        col1, col2 = st.columns([0.1, 5])
        with col1:
            arrow = "▼" if st.session_state.show_consistency else "▶"
            if st.button(arrow, key="consistency_arrow", help="折叠/展开"):
                st.session_state.show_consistency = not st.session_state.show_consistency
        with col2:
            st.markdown("<h2 style='font-size: 28px; margin: 0;'>上下文一致性检查</h2>", unsafe_allow_html=True)
        if st.session_state.show_consistency:
            for res in st.session_state.results:
                if "✅" in res["状态"]:
                    st.success(f"**{res['项目']}**：{res['状态']}")
                else:
                    st.error(f"**{res['项目']}**：{res['状态']}")
                with st.expander("查看详情"):
                    st.text(res["显示值"])

    # 下载审核报告
    if st.session_state.get("negative_issues") or st.session_state.get("issues") or st.session_state.get("results"):
        sections = []
        if st.session_state.get("negative_issues"):
            sections.append((st.session_state.get("negative_issues"), "负面清单检查结果"))
        if st.session_state.get("issues"):
            sections.append((st.session_state.get("issues"), "内容合理性检查结果"))
        if st.session_state.get("results"):
            sections.append((st.session_state.get("results"), "上下文一致性检查结果"))
        
        doc_bytes = get_download_button_data(
            sections=sections,
            title="招标文件审查报告",
            filename="招标文件审查报告.docx"
        )
        
        st.download_button(
            label="📄 下载审核报告",
            data=doc_bytes,
            file_name="招标文件审查报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("点击「开始检查」后，即可下载审核报告")
else:  # 响应文件
    st.info("响应文件功能正在开发中，敬请期待...")