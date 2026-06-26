# ============================================
# report_generator.py - 统一报告生成模块
# 可被多个 Streamlit 应用复用
# ============================================

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from io import BytesIO
import re


def generate_report(
    sections=None,
    title="审查报告",
    filename="审查报告.docx"
):
    """
    生成审查报告 Word 文档
    
    参数:
        sections: 章节数据列表，每个元素为 (数据列表, 章节标题)
                 例如: [(negative_issues, "负面清单检查结果"), (issues, "内容合理性检查结果")]
                 也可以传入字典: {"负面清单检查结果": negative_issues, "内容合理性检查结果": issues}
        title: 报告大标题
        filename: 下载文件名
    
    返回:
        BytesIO: Word 文档的二进制数据
    """
    
    # 初始化文档
    doc = Document()
    
    # ---------- 处理参数：支持多种传入方式 ----------
    if sections is None:
        sections = []
    
    # 如果传入的是字典，转换为列表
    if isinstance(sections, dict):
        sections = [(data, title) for title, data in sections.items() if data]
    
    # 如果传入的是单个章节，转换为列表
    if isinstance(sections, tuple) and len(sections) == 2:
        sections = [sections]
    
    # 如果传入的是列表但元素不是元组，尝试转换
    if isinstance(sections, list) and sections and not isinstance(sections[0], tuple):
        # 假设是 [data1, data2, ...] 但没有标题
        # 这种情况需要调用方传入正确的格式
        pass
    
    # ---------- 过滤掉空数据 ----------
    valid_sections = []
    for data, section_title in sections:
        if data and len(data) > 0:
            valid_sections.append((data, section_title))
    
    if not valid_sections:
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    
    # ---------- 清理 HTML 标签 ----------
    def clean_html(text):
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    # ---------- 辅助函数：添加空行 ----------
    def add_blank_lines(count=1):
        for _ in range(count):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # ---------- 辅助函数：判断标题是否超过一行 ----------
    def is_multiline_title(title_text):
        estimated_width = 0
        for ch in title_text:
            if '\u4e00' <= ch <= '\u9fff':
                estimated_width += 1
            else:
                estimated_width += 0.5
        return estimated_width / 35 > 1
    
    # ---------- 辅助函数：添加带颜色的标题行 ----------
    def add_colored_title(title_text, is_violation):
        p = doc.add_paragraph()
        if is_multiline_title(title_text):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(12)
        if is_violation:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(0, 128, 0)
        return p
    
    # ---------- 辅助函数：拆分违规内容 ----------
    def split_violations(text):
        if not text:
            return [text]
        text = clean_html(text)
        parts = re.split(r'(?=【违规\d+[ -][^】]*】)', text)
        result = [p.strip() for p in parts if p.strip()]
        if len(result) <= 1:
            parts = re.split(r'(?=【违规\d+】)', text)
            result = [p.strip() for p in parts if p.strip()]
        if len(result) <= 1:
            parts = re.split(r'(?=违规\d+[ -][^，,。]*[，,。]?)', text)
            result = [p.strip() for p in parts if p.strip() and '违规' in p]
        if len(result) <= 1:
            if '<br>' in text or '<br />' in text:
                parts = re.split(r'<br\s*/?>', text)
                result = [p.strip() for p in parts if p.strip()]
        if len(result) <= 1:
            return [text]
        cleaned = [r.strip() for r in result if r.strip()]
        return cleaned if cleaned else [text]
    
    # ---------- 辅助函数：添加原文段落 ----------
    def add_original_text(original_text):
        violations = split_violations(original_text)
        if len(violations) > 1:
            for i, violation in enumerate(violations):
                if violation:
                    if not re.search(r'【违规\d+', violation):
                        prefix = f"【违规{i+1}】"
                        p = doc.add_paragraph(f"原文：{prefix}{violation}")
                    else:
                        p = doc.add_paragraph(f"原文：{violation}")
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            p = doc.add_paragraph(f"原文：{original_text}")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # ---------- 辅助函数：添加标签段落 ----------
    def add_labeled_paragraph(label, content):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0.5)
        run_label = para.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_content = para.add_run(content)
        run_content.font.size = Pt(12)
        return para
    
    # ---------- 辅助函数：判断数据类型并渲染 ----------
    def render_items(data, section_title):
        """根据数据内容自动识别类型并渲染"""
        if not data or len(data) == 0:
            return
        
        first_item = data[0]
        
        # 检测数据类型
        if "规则" in first_item and "原文" in first_item:
            # 负面清单或内容合理性类型
            for item_idx, item in enumerate(data):
                is_violation = not item.get("通过", True)
                title_text = f"【{item.get('规则', '')}】 {item.get('详情', '')}"
                add_colored_title(title_text, is_violation)
                original_text = item.get('原文', '')
                add_original_text(original_text)
                if item_idx < len(data) - 1:
                    add_blank_lines(1)
        
        elif "项目" in first_item and "状态" in first_item:
            # 上下文一致性类型
            for item_idx, item in enumerate(data):
                is_violation = "❌" in item.get("状态", "") or "不一致" in item.get("状态", "")
                title_text = f"【{item.get('项目', '')}】 {item.get('状态', '')}"
                add_colored_title(title_text, is_violation)
                
                display_value = item.get('显示值', '')
                
                gonggao_text = ""
                xuzhi_text = ""
                bijiao_text = ""
                
                lines = display_value.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('公告:') or line.startswith('公告：'):
                        gonggao_text = line
                    elif line.startswith('须知:') or line.startswith('须知：'):
                        xuzhi_text = line
                    elif line.startswith('比较依据:') or line.startswith('比较依据：'):
                        bijiao_text = line
                
                if not gonggao_text and not xuzhi_text and not bijiao_text:
                    parts = display_value.split('  ')
                    for part in parts:
                        part = part.strip()
                        if part.startswith('公告:') or part.startswith('公告：'):
                            gonggao_text = part
                        elif part.startswith('须知:') or part.startswith('须知：'):
                            xuzhi_text = part
                        elif part.startswith('比较依据:') or part.startswith('比较依据：'):
                            bijiao_text = part
                
                if not gonggao_text and not xuzhi_text and not bijiao_text:
                    match = re.search(r'公告[：:]\s*([^，,。\n]*?(?:[，,。][^，,。]*?)*?)(?=\s*须知[：:]|\s*比较依据[：:]|$)', display_value)
                    if match and match.group(1).strip():
                        gonggao_text = f"公告：{match.group(1).strip()}"
                    match = re.search(r'须知[：:]\s*([^，,。\n]*?(?:[，,。][^，,。]*?)*?)(?=\s*比较依据[：:]|$)', display_value)
                    if match and match.group(1).strip():
                        xuzhi_text = f"须知：{match.group(1).strip()}"
                    match = re.search(r'比较依据[：:]\s*(.+?)$', display_value)
                    if match and match.group(1).strip():
                        bijiao_text = f"比较依据：{match.group(1).strip()}"
                
                if gonggao_text:
                    if '：' in gonggao_text:
                        label, content = gonggao_text.split('：', 1)
                        add_labeled_paragraph(f"{label}：", content)
                    elif ':' in gonggao_text:
                        label, content = gonggao_text.split(':', 1)
                        add_labeled_paragraph(f"{label}:", content)
                    else:
                        p = doc.add_paragraph(gonggao_text)
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_after = Pt(0)
                else:
                    p = doc.add_paragraph("公告：（未提取到）")
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(0)
                
                if xuzhi_text:
                    if '：' in xuzhi_text:
                        label, content = xuzhi_text.split('：', 1)
                        add_labeled_paragraph(f"{label}：", content)
                    elif ':' in xuzhi_text:
                        label, content = xuzhi_text.split(':', 1)
                        add_labeled_paragraph(f"{label}:", content)
                    else:
                        p = doc.add_paragraph(xuzhi_text)
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_after = Pt(0)
                else:
                    p = doc.add_paragraph("须知：（未提取到）")
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(0)
                
                if bijiao_text:
                    if '：' in bijiao_text:
                        label, content = bijiao_text.split('：', 1)
                        add_labeled_paragraph(f"{label}：", content)
                    elif ':' in bijiao_text:
                        label, content = bijiao_text.split(':', 1)
                        add_labeled_paragraph(f"{label}:", content)
                    else:
                        p = doc.add_paragraph(bijiao_text)
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_after = Pt(0)
                else:
                    p = doc.add_paragraph("比较依据：无比较依据")
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(0)
                
                if item_idx < len(data) - 1:
                    add_blank_lines(1)
        
        else:
            # 未知类型：直接显示原始数据
            for item_idx, item in enumerate(data):
                p = doc.add_paragraph(str(item))
                p.paragraph_format.left_indent = Cm(0.5)
                if item_idx < len(data) - 1:
                    add_blank_lines(1)
    
    # ---------- 设置文档默认字体和行距 ----------
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    
    # ---------- 大标题 ----------
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(22)
    heading.runs[0].font.bold = True
    
    # ---------- 过滤掉空数据 ----------
    valid_sections = []
    for data, section_title in sections:
        if data and len(data) > 0:
            valid_sections.append((data, section_title))
    
    if not valid_sections:
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    
    # ---------- 生成各章节 ----------
    chinese_numbers = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    
    for idx, (data, section_title) in enumerate(valid_sections):
        if idx > 0:
            add_blank_lines(1)
        
        # 章节编号
        number = chinese_numbers[idx] if idx < len(chinese_numbers) else str(idx + 1)
        doc.add_heading(f"{number}、{section_title}", level=1)
        doc.paragraphs[-1].runs[0].font.size = Pt(16)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
        
        # 渲染数据
        render_items(data, section_title)
    
    # ---------- 保存到 BytesIO ----------
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


def get_download_button_data(
    sections=None,
    title="审查报告",
    filename="审查报告.docx"
):
    """
    生成报告并返回 BytesIO 数据
    
    用法:
        # 方式1：传入列表
        doc_bytes = get_download_button_data(
            sections=[
                (negative_issues, "负面清单检查结果"),
                (issues, "内容合理性检查结果"),
                (results, "上下文一致性检查结果"),
                (other_data, "其他检查结果")  # 任意类型都可以
            ],
            title="招标文件审查报告"
        )
        
        # 方式2：传入字典
        doc_bytes = get_download_button_data(
            sections={
                "负面清单检查结果": negative_issues,
                "内容合理性检查结果": issues,
                "上下文一致性检查结果": results,
                "其他检查结果": other_data
            },
            title="招标文件审查报告"
        )
    """
    return generate_report(
        sections=sections,
        title=title,
        filename=filename
    )